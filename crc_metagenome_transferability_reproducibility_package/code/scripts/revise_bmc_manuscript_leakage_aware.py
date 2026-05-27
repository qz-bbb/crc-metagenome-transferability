from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


RUN_DIR = Path(r"D:\1\projects\gut_microbiome_colorectal_cancer\runs\run_20260525_182306")
BMC = RUN_DIR / "BMC_submission"
MANUSCRIPT_DOCX = BMC / "manuscript" / "manuscript_BMC_Microbiology.docx"
MANUSCRIPT_MD = BMC / "manuscript" / "manuscript_BMC_Microbiology.md"
COVER_DOCX = BMC / "cover_letter" / "cover_letter_BMC_Microbiology.docx"
COVER_MD = BMC / "cover_letter" / "cover_letter_BMC_Microbiology.md"

TITLE = "Transferability-aware prioritization of colorectal cancer gut metagenome signals across public cohorts"
AUTHOR = "Zhun Qiu"
AFFILIATION = "Northeast Forestry University, Harbin, Heilongjiang, China"
EMAIL = "qz@nefu.edu.cn"

SPECIES_NAMES = [
    "Parvimonas micra",
    "Gemella morbillorum",
    "Peptostreptococcus stomatis",
    "Roseburia",
    "Faecalibacterium",
]


def get_existing_references() -> list[str]:
    text = MANUSCRIPT_MD.read_text(encoding="utf-8", errors="ignore")
    if "## References" not in text:
        return []
    refs = text.split("## References", 1)[1].split("## Figure legends", 1)[0].strip()
    return [p.strip().replace("\n", " ") for p in re.split(r"\n\s*\n", refs) if p.strip()]


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def enable_line_numbers(section) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "newPage")
    sect_pr.append(ln)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    enable_line_numbers(section)
    add_page_number(section)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)
    for name, size in [("Title", 16), ("Heading 1", 14), ("Heading 2", 12)]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.paragraph_format.line_spacing = 2


def add_para(doc: Document, text: str = "", style: str | None = None, bold_label: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_label and text.startswith(bold_label):
        run = p.add_run(bold_label)
        run.bold = True
        text = text[len(bold_label) :]
    start = 0
    while start < len(text):
        matches = [(text.find(name, start), name) for name in SPECIES_NAMES if text.find(name, start) >= 0]
        if not matches:
            p.add_run(text[start:])
            break
        pos, name = min(matches)
        if pos > start:
            p.add_run(text[start:pos])
        run = p.add_run(name)
        run.italic = True
        start = pos + len(name)


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def build_manuscript() -> None:
    references = get_existing_references()
    doc = Document()
    style_document(doc)
    doc.add_paragraph(TITLE, style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, AUTHOR)
    add_para(doc, f"Affiliation: {AFFILIATION}")
    add_para(doc, f"Correspondence: {AUTHOR}, {EMAIL}")

    h1(doc, "Abstract")
    add_para(doc, "Background: Public processed metagenomic resources enable sample-level reanalysis of colorectal cancer (CRC)-associated gut microbiome signals across cohorts without starting from raw sequencing files. We tested whether association-ranked microbial signals, stability-ranked signals, and transferability-aware panels behave differently under cohort shift.", bold_label="Background:")
    add_para(doc, "Results: The analysis included 1,395 CRC/control samples across 11 study labels and retained 301 species-level features. Complete-case Aitchison-space variance partitioning attributed more retained CLR variation to study label than to CRC status (term-deletion partial R² 8.02% versus 0.51%; 999 permutations, p ≤ 0.001 for both). A study-adjusted screen identified 124 species at FDR < 0.10, while 29 candidates met a ≥60% same-direction rule and 18 met high-consistency criteria. In leakage-aware LOSO benchmarking, training-only top29-q and training-only ≥60% stability panels converged in all folds (median AUROC 0.782), whereas the stricter high-consistency panel was smaller and did not outperform top-q ranking (median paired AUROC difference -0.017; bootstrap 95% CI -0.038 to 0.004). Retrospective fixed-panel transfer showed that full-data top29-q had the highest off-diagonal AUROC, while high-consistency candidates had lower transferability loss. Taxonomy-defined guild scores provided conservative ecological context but not functional or source inference.", bold_label="Results:")
    add_para(doc, "Conclusions: Association strength, direction stability, ecological coherence, and cohort portability were not interchangeable in public CRC gut metagenomes. Transferability-aware prioritization is useful as a conservative interpretation framework, but these processed-data analyses do not establish causal, clinical diagnostic, stage-specific, mechanistic, functional, strain-level, oral-source, or raw-read reprocessing claims.", bold_label="Conclusions:")

    h1(doc, "Keywords")
    add_para(doc, "Colorectal cancer; gut microbiome; metagenomics; curatedMetagenomicData; compositional data analysis; cohort-aware analysis; transferability; reproducibility")

    h1(doc, "Background")
    add_para(doc, "Colorectal cancer (CRC)-associated gut microbiome differences have been reported across several fecal metagenomic studies, including species-level signals involving oral-associated taxa, butyrate-associated commensals, and other community shifts [1-7]. These studies have helped establish that CRC status can be associated with reproducible microbial differences, yet they also show that effect sizes, taxonomic resolution, and cohort portability can vary across populations and study designs.")
    add_para(doc, "Public harmonized resources such as curatedMetagenomicData and ExperimentHub-derived processed tables make it possible to work directly with sample-level abundance matrices rather than published summary statistics [9]. Related resources such as GMrepo and MicrobiomeHD provide additional disease-microbiome context, but they differ in schema, taxonomic resolution, and intended use [8,10,11]. For CRC, these resources create an opportunity to evaluate how far processed matrices can support reproducible candidate prioritization without requiring raw-read reprocessing at the first stage.")
    add_para(doc, "A central challenge is that study-adjusted association strength does not necessarily imply portability across cohorts. A species may be strongly associated with CRC status in a pooled processed-matrix model but show weaker direction stability, high heterogeneity, or limited transfer when one cohort is held out. Conversely, a more modest association may be more directionally stable and therefore more useful for transparent candidate prioritization. These layers should be evaluated separately rather than collapsed into one significance threshold or one model performance number.")
    add_para(doc, "We tested whether association-ranked, robustness-ranked, and transferability-aware microbial signals identify the same CRC-associated candidates in public processed gut metagenomes. We further evaluated whether study structure explains a substantial component of Aitchison-space variation and whether transferability-aware panel prioritization provides a more conservative and interpretable candidate set than pooled association ranking alone. The analysis was restricted to sample-level processed abundance matrices and was not designed to provide raw-read reprocessing, clinical diagnostic validation, causal inference, stage-specific biomarkers, functional profiling, or strain-level claims.")

    h1(doc, "Methods")
    h2(doc, "Data source and feature filtering")
    add_para(doc, "The analysis used curatedMetagenomicData/ExperimentHub-derived processed CRC/control resources with sample metadata, a species-level percent relative-abundance matrix, and a species feature index [9]. Across study-level inputs, species absent from a study-level matrix were filled as zero in the cross-study union matrix. The analysis-ready matrix contained 944 species-level features and 1,395 samples across 11 study labels.")
    add_para(doc, "The primary feature filter retained species with overall prevalence ≥5% and presence in at least three study labels. Prevalence was calculated from non-zero relative abundance. This filter retained 301 of 944 species-level features. Centered log-ratio (CLR) transformation was performed after filtering. The pseudocount was defined as one-half of the minimum positive retained abundance value, equal to 2.5e-05 in percent relative-abundance units.")
    h2(doc, "Study-adjusted association and covariate sensitivity")
    add_para(doc, "For each retained species, the primary association model was an ordinary least-squares model with study indicators: CLR_taxon ~ CRC_status + study. CRC_status was encoded as CRC = 1 and control = 0, so positive coefficients indicate CRC-higher CLR abundance. Heteroscedasticity-consistent HC3 standard errors were used, and P values were adjusted across tested features with the Benjamini-Hochberg false discovery rate (FDR) procedure [13].")
    add_para(doc, "The primary covariate sensitivity model used complete-case samples with age, BMI, and gender metadata: CLR_taxon ~ CRC_status + study + age_z + BMI_z + gender. Age and BMI were z-standardized within the analysis dataset. Country and sequencing platform metadata were summarized as context, but the primary robustness interpretation used the study-adjusted and study + age + BMI + gender models to avoid over-interpreting sparse design strata.")
    h2(doc, "Aitchison ordination and variance partitioning")
    add_para(doc, "Aitchison PCA was computed from the filtered CLR matrix. CRC/control compositional separation after study adjustment was also evaluated by residualizing sample CLR vectors against study-specific feature means and permuting CRC/control labels within study labels.")
    add_para(doc, "To quantify the relative contribution of study label, CRC/control status, and available covariates, retained CLR profiles were analyzed with distance-equivalent multivariate linear models. Euclidean distance on retained CLR profiles was treated as Aitchison distance after CLR transformation. Complete-case variance partitioning used the formula retained_CLR ~ study_label + CRC_status + age_z + BMI_z + gender. Sequential R² was calculated in that order. Term-deletion partial R² was calculated by comparing the full model with reduced models omitting each term while retaining the other listed terms. Term-level permutation P values used 999 permutations; because 999 permutations set a lower resolution of 0.001, values at this boundary are reported as p ≤ 0.001.")
    h2(doc, "Direction stability and heterogeneity")
    add_para(doc, "Cross-cohort stability was evaluated using per-study effect directions for each retained species. A stable candidate was defined before interpretation as a species with study-adjusted BH FDR < 0.10 and same-direction support in at least 60% of eligible study labels. Per-cohort effect estimates were also summarized with DerSimonian-Laird random-effects analysis, including random-effects estimates, random-effects q values, τ², and I². This heterogeneity analysis was treated as an audit layer for processed-data consistency rather than as external validation.")
    h2(doc, "Leakage-aware LOSO panel benchmarking")
    add_para(doc, "Primary panel benchmarking was performed in a leakage-aware LOSO design. For each held-out study label, feature selection used only the remaining training studies. The training-only association screen refit CLR_taxon ~ CRC_status + study across training studies and ranked taxa by training-only BH q value.")
    add_para(doc, "Four deterministic training-fold panels were evaluated: all 301 retained species, the top 29 training-only q-ranked species, a training-only stability panel requiring FDR < 0.10 and ≥60% same-direction support across eligible training cohorts, and a stricter high-consistency panel requiring FDR < 0.10 and ≥80% same-direction support. The stability panel was capped at 29 taxa and the high-consistency panel at 18 taxa; if fewer taxa met a rule, all eligible taxa were retained and the actual panel size was reported.")
    add_para(doc, "For each fold, scaling and elastic-net model fitting were restricted to training studies and then applied to the held-out study. The logistic regression model used saga optimization, elastic-net penalty with l1_ratio = 0.5, C = 0.5, class_weight = balanced, max_iter = 2000, and random_state = 42. Random 29-species panels were evaluated 500 times per held-out fold, including unmatched random panels and prevalence-matched random panels where feasible. These analyses were separability stress tests, not clinical model development.")
    h2(doc, "Retrospective fixed-panel transferability stress testing")
    add_para(doc, "Secondary fixed-panel analyses retained the full-data panels from the previous prioritization layer: all 301 retained species, the full-data top 29 species ranked by study-adjusted q value, the full-data 29 stable candidates, the full-data 18 high-consistency candidates, and random 29-species panels. These fixed panels were evaluated with LOSO and pairwise train-study × test-study transfer, but they were labeled as retrospective fixed-panel transferability stress tests because the feature sets were selected from full-data results. Transferability loss was defined as median within-study pairwise AUROC minus median off-diagonal pairwise AUROC.")
    h2(doc, "Ecological guild score analyses")
    add_para(doc, "Three transparent taxonomy-defined guild panels were constructed from taxa present in the retained processed species matrix: a butyrate/SCFA-associated commensal panel, an oral/pathobiont-associated panel, and a Bacteroides/Enterobacteriaceae context panel. Guild scores were defined as the sample-wise mean CLR value across retained taxa in each panel. Orientation-coded scores were used only for score-only separability summaries and were not treated as functional measurements.")
    add_para(doc, "For each guild score, analyses included study-adjusted association, age/BMI/gender covariate sensitivity, per-cohort Hedges g, DerSimonian-Laird random-effects analysis, leave-one-cohort-out robustness, and optional LOSO score-only AUROC. Guild membership is reported exactly in Supplementary Table S9. These guild scores are ecological annotation summaries; they are not functional profiles, strain-level analyses, mechanistic assays, or oral-source attribution models.")
    h2(doc, "Panel comparison and claim decisions")
    add_para(doc, "LOSO panel comparisons used paired fold-level AUROC differences across the 11 held-out studies. Median and mean paired differences were reported with bootstrap 95% confidence intervals, and Wilcoxon signed-rank P values were treated as exploratory because the number of folds was small. Pairwise transfer comparisons used descriptive off-diagonal AUROC distributions with bootstrap confidence intervals and avoided strong significance claims because train-study × test-study pairs are dependent. A claim decision table was generated to separate supported wording from prohibited causal, clinical, functional, strain-level, oral-source, and raw-read claims.")
    h2(doc, "Robustness analyses")
    add_para(doc, "We further evaluated whether the association workflow was sensitive to feature-filtering thresholds, pseudocount choice, and individual-cohort removal. Feature-filter sensitivity used prevalence thresholds of 3%, 5%, and 10% crossed with cohort-presence thresholds of two, three, or four study labels. Pseudocount sensitivity used one-quarter, one-half, and one times the minimum positive retained abundance. Leave-one-cohort-out association sensitivity refit the association screen after removing each study label.")
    h2(doc, "Use of AI-assisted tools")
    add_para(doc, "AI-assisted coding and language tools were used to support code refactoring, manuscript editing, and formatting during manuscript preparation. The author reviewed and verified all tool-assisted outputs, reran or checked the analyses where applicable, and takes full responsibility for the content, results, interpretation, and submission of the manuscript.")

    h1(doc, "Results")
    h2(doc, "Study structure motivates cohort-aware transferability analysis")
    add_para(doc, "The analysis covered 1,395 samples, including 701 CRC and 694 control records across 11 study labels. After prevalence and cohort-presence filtering, 301 of 944 species-level features were retained for CLR-based analysis.")
    add_para(doc, "Complete-case Aitchison-space variance partitioning used 1,380 samples. Study label explained more retained CLR variation than CRC/control status: term-deletion partial R² was 8.02% for study label and 0.51% for CRC status after accounting for the other listed terms, with 999-permutation p ≤ 0.001 for both. Age, gender, and BMI explained smaller fractions (0.37%, 0.21%, and 0.15%, respectively), and the full model left 90.2% residual variation. Aitchison PCA showed broad CRC/control overlap, while the within-study residualized permutation test detected CRC/control compositional separation (pseudo-F 7.90, p = 0.0002, 4,999 permutations; Figure 1; Supplementary Table S6).")
    h2(doc, "Study-adjusted association identifies a broad CRC-associated species landscape")
    add_para(doc, "The study-adjusted CLR association screen tested 301 species and identified 124 features at FDR < 0.10 and 107 features at FDR < 0.05. The strongest CRC-higher associations included Parvimonas micra, Gemella morbillorum, and Peptostreptococcus stomatis. Positive coefficients indicate CRC-higher CLR abundance and negative coefficients indicate control-higher abundance (Figure 2).")
    add_para(doc, "The complete-case covariate model used 1,380 samples with age, BMI, and gender metadata. This sensitivity analysis identified 137 features at FDR < 0.10, preserved all 29 primary stable candidates, showed 91/100 overlap among the top baseline features, and had all-feature direction concordance of 0.947. These results indicate that the main association landscape was not driven solely by the available age, BMI, or gender metadata fields.")
    h2(doc, "Direction stability restricts association signals to a smaller candidate set")
    add_para(doc, "Cross-cohort stability was more restrictive than study-adjusted association strength. Among the 124 FDR < 0.10 features, 29 candidates also met the ≥60% same-direction cohort-support rule. Eighteen stable candidates were classified as high-consistency candidates after integrating leave-one-cohort and random-effects support.")
    add_para(doc, "Random-effects heterogeneity auditing tested all 301 retained species. Ninety-four species had random-effects FDR < 0.10, the median I² was 40.6%, and 17 species had I² ≥75%. These results show that study-adjusted association, cross-cohort direction stability, and random-effects support select overlapping but non-identical signals (Figure 3).")
    h2(doc, "Training-only panel benchmarking separates association ranking from stricter prioritization")
    add_para(doc, "Leakage-aware LOSO benchmarking recomputed feature selection within each training set before held-out evaluation. The all-301 panel had median held-out AUROC 0.752. The training-only top29-q panel had median AUROC 0.782, minimum AUROC 0.582, and median average precision 0.834. The training-only ≥60% stability rule selected the same 29 taxa as the training-only top29-q panel in all folds, producing identical held-out AUROC values. The stricter high-consistency panel retained 18 taxa per fold and had median AUROC 0.770.")
    add_para(doc, "Fold-level paired comparisons did not support a claim that transferability-aware panels outperformed pooled-q ranking. The high-consistency panel minus top29-q median AUROC difference was -0.017, with bootstrap 95% CI -0.038 to 0.004 and exploratory Wilcoxon p = 0.067. The top29-q panel minus all-301 median AUROC difference was 0.009, with bootstrap 95% CI -0.032 to 0.045. Random-panel comparisons provided context but did not justify clinical performance claims: training-only stability AUROC exceeded prevalence-matched random panels in most folds, but unmatched random panels were closer and cohort-dependent (Figure 7; Supplementary Tables S7 and S10).")
    h2(doc, "Fixed-panel pairwise transfer exposes cohort-dependent portability")
    add_para(doc, "Retrospective fixed-panel transferability stress testing retained full-data panels but labeled them separately because feature selection used the full processed matrix. In this secondary analysis, the full-data top29-q panel had the highest LOSO median AUROC (0.834) and off-diagonal pairwise median AUROC (0.733). The full-data stable29 panel had lower separability (LOSO median AUROC 0.673; off-diagonal median AUROC 0.595), and the high-consistency18 panel had LOSO median AUROC 0.693 and off-diagonal median AUROC 0.634.")
    add_para(doc, "Pairwise train-study × test-study transfer remained cohort-dependent. The high-consistency18 panel had lower transferability loss (0.186) than all 301 features (0.345) and stable29 (0.308), and slightly lower loss than top29-q (0.205), but it did not maximize off-diagonal AUROC. These results support a divergence framing: association-ranked panels can maximize separability, while stricter prioritization can produce a smaller and more interpretable panel with a different portability profile (Figures 4-5; Supplementary Table S8).")
    h2(doc, "Ecological guild scores provide conservative biological interpretation")
    add_para(doc, "Taxonomy-defined ecological guild scores provided interpretable but bounded context. The butyrate/SCFA-associated commensal panel contained 26 retained taxa, including 10 stable candidates, and was control-higher in both study-adjusted association and random-effects Hedges g analysis (random-effects Hedges g -0.466, q = 2.8e-05, I² = 67.6%).")
    add_para(doc, "The oral/pathobiont-associated panel contained 28 retained taxa but only one stable candidate; it was CRC-higher at the guild-score level (random-effects Hedges g 0.262, q = 0.0017, I² = 45.7%). The Bacteroides/Enterobacteriaceae context panel contained nine retained taxa and three stable candidates, with random-effects Hedges g 0.193, q = 0.018, and I² = 47.2%. Score-only LOSO AUROC medians were modest: 0.614 for butyrate/SCFA, 0.558 for oral/pathobiont, and 0.552 for Bacteroides/Enterobacteriaceae. These results support ecological annotation axes, not functional activity, oral-source attribution, or mechanism claims (Figure 8; Supplementary Table S9).")
    h2(doc, "Sensitivity analyses support robustness but not clinical validation")
    add_para(doc, "Feature-filter sensitivity produced 94-145 FDR < 0.10 features across filtering configurations, with top-100 overlap ranging from 78 to 100 and stable-candidate counts ranging from 28 to 29. Pseudocount sensitivity produced 124-127 FDR < 0.10 features, top-100 overlap of at least 100/100, and minimum direction concordance of 0.993. Stricter stability thresholds retained 20 candidates at ≥70% same-direction support and 11 candidates at ≥80% same-direction support.")
    add_para(doc, "LOCO association sensitivity supported 26 of the 29 stable candidates under the predefined robustness rule. Together, these sensitivity analyses show that the main processed-data association landscape was robust to several analytic choices, while candidate interpretation remained dependent on the chosen support layer (Figure 6; Supplementary Figure S2; Supplementary Table S4).")

    h1(doc, "Discussion")
    add_para(doc, "The main contribution of this study is not another CRC microbial signature, but a transferability-aware prioritization framework showing that association strength, stability, ecological coherence, and portability do not fully coincide. The leakage-aware analysis was especially important: once feature selection was restricted to training studies, the ≥60% stability panel converged with the training-only top-q panel, and the stricter high-consistency panel became more conservative but did not improve AUROC. This supports a claim of divergence between association ranking and portability, not a claim of superior predictive performance.")
    add_para(doc, "Variance partitioning provides the rationale for this cohort-aware framing. Study label explained substantially more retained CLR variation than CRC/control status, while CRC status remained detectable as a small cohort-embedded signal. This pattern is expected in public processed metagenome matrices assembled across countries, platforms, study designs, and case-control definitions. It means that candidate prioritization should be evaluated under cohort shift rather than inferred from pooled association strength alone.")
    add_para(doc, "The panel benchmarking results sharpen the interpretation. Full-data fixed-panel comparisons are useful retrospective stress tests, but they can look optimistic if treated as held-out validation after global feature selection. The training-only LOSO analysis reduces that concern and shows that apparent separability depends on the selection rule. The q-ranked panel was not displaced by the stability panel in leakage-aware folds; the more restrictive high-consistency panel should therefore be presented as an interpretable shortlist, not a higher-performing classifier.")
    add_para(doc, "The ecological guild scores add conservative biological context. Butyrate/SCFA-associated commensal taxa tended to be control-higher, while oral/pathobiont-associated taxa were CRC-higher at the guild-score level. However, these panels were defined from taxonomic names in a species-level processed matrix. They should not be interpreted as direct evidence of metabolite production, host pathway activity, oral transmission, strain behavior, or functional pathway activity.")
    add_para(doc, "The limitations remain important. The analysis used public processed tables only; no raw reads were reprocessed; no disease-stage model, functional profile, strain-level profile, independent patient-level validation matrix, or diagnostic calibration was analyzed. The results support processed-data association, robustness, heterogeneity, and transferability claims, but not causal, clinical diagnostic, stage-specific, mechanistic, functional, strain-level, oral-source, or raw-read reprocessing claims.")

    h1(doc, "Conclusions")
    add_para(doc, "Public processed CRC/control gut metagenome matrices can support transferability-aware candidate prioritization when association strength, direction stability, ecological coherence, heterogeneity, and cohort portability are evaluated separately. In this 1,395-sample reanalysis, study structure explained more Aitchison-space variation than CRC/control status, and leakage-aware panel benchmarking showed that training-only top-q and ≥60% stability panels converged while stricter high-consistency prioritization produced a smaller, more conservative shortlist.")
    add_para(doc, "The central result is divergence, not diagnostic improvement. Association-ranked panels maximized separability in several retrospective comparisons, whereas transferability-aware and ecological guild analyses clarified which signals were more stable, interpretable, or cohort-dependent. These findings support a reproducible processed-data candidate-prioritization workflow under cohort shift, but they do not establish causal, clinical diagnostic, stage-specific, mechanistic, functional, strain-level, oral-source, or raw-read claims.")

    h1(doc, "List of abbreviations")
    for item in [
        "AUROC: area under the receiver operating characteristic curve",
        "BH: Benjamini-Hochberg",
        "CLR: centered log-ratio",
        "CRC: colorectal cancer",
        "FDR: false discovery rate",
        "LOCO: leave-one-cohort-out",
        "LOSO: leave-one-study-out",
        "PCA: principal component analysis",
    ]:
        add_para(doc, item)

    h1(doc, "Declarations")
    for heading, text in [
        ("Ethics approval and consent to participate", "This study is a secondary analysis of publicly available, de-identified processed microbiome abundance tables and associated metadata. No new human participants were recruited, no intervention was performed, and no identifiable private information was accessed. Ethics approval and consent to participate were therefore not required for this secondary analysis."),
        ("Consent for publication", "Not applicable."),
        ("Availability of data and materials", "The public source data analyzed in this study are available through curatedMetagenomicData/ExperimentHub-derived resources. The analysis-ready metadata identifiers, species-level abundance processing scripts, statistical analysis scripts, figure-generation scripts, software environment files, random seeds, and complete result tables generated for this study are available at GitHub: [ADD GITHUB REPOSITORY LINK]. An archived version of the analysis code and result tables is available at Zenodo: [ADD ZENODO DOI]. Supplementary tables are also provided as machine-readable additional files."),
        ("Competing interests", "The author declares that he has no competing interests."),
        ("Funding", "The author received no specific funding for this work."),
        ("Authors' contributions", "ZQ conceived the study, curated the public processed datasets, designed and performed the analyses, generated the figures and tables, interpreted the results, wrote and revised the manuscript, and approved the final manuscript."),
        ("Acknowledgements", "Not applicable."),
        ("Authors' information", "Not applicable."),
    ]:
        h2(doc, heading)
        add_para(doc, text)

    h1(doc, "Additional files")
    additional = [
        "Additional file 1: Supplementary Figure S1. Covariate completeness and primary sensitivity comparison.",
        "Additional file 2: Supplementary Figure S2. Stable-candidate evidence map.",
        "Additional file 3: Supplementary Table S1. Cohort and sample characteristics.",
        "Additional file 4: Supplementary Table S2. Stable candidate taxa.",
        "Additional file 5: Supplementary Table S3. Leave-one-study-out performance.",
        "Additional file 6: Supplementary Table S4. Stable-candidate evidence map.",
        "Additional file 7: Supplementary Table S5. Literature-prior oral-associated genus panel.",
        "Additional file 8: Supplementary Table S6. Aitchison variance partitioning.",
        "Additional file 9: Supplementary Table S7. Compact leakage-aware LOSO panel benchmarking summary.",
        "Additional file 10: Supplementary Table S8. Retrospective fixed global panel benchmarking.",
        "Additional file 11: Supplementary Table S9. Ecological guild membership and score results.",
        "Additional file 12: Supplementary Table S10. Panel comparison statistics.",
        "Additional file 13: Supplementary Table S11. Claim decision table.",
    ]
    for item in additional:
        add_para(doc, item)

    h1(doc, "References")
    for ref in references:
        add_para(doc, ref)

    h1(doc, "Figure legends")
    legends = [
        "Figure 1. Study structure and retained feature space. (A) Aitchison PCA of retained species-level CLR profiles. (B) Term-deletion partial variation by study label, CRC status, and available covariates. (C) Feature prevalence-abundance landscape showing retained and filtered species.",
        "Figure 2. Study-adjusted species associations. (A) Association volcano plot for 301 retained species. (B) Coefficient intervals for top candidates. (C) Raw-abundance sensitivity comparison. (D) Baseline versus age, BMI, and gender sensitivity model comparison. Positive coefficients indicate CRC-higher CLR abundance.",
        "Figure 3. Cross-cohort stability and heterogeneity. (A) Random-effects summaries for candidate taxa. (B) Relationship between association strength and heterogeneity. (C) Per-cohort effect heatmap for selected candidates.",
        "Figure 4. LOSO separability stress test and transportability ranking. (A) Held-out-study AUROC values with bootstrap 95% confidence intervals. (B) Average precision across held-out studies. (C) Held-out predicted CRC probability distributions by true condition. (D) Exploratory composite transportability ranking. These panels describe processed-table separability and candidate ordering, not clinical validation.",
        "Figure 5. Fixed-panel and pairwise study transfer. (A) Retrospective fixed-panel LOSO summaries. (B) Pairwise train-study × test-study AUROC matrix. (C) Off-diagonal transfer distribution. (D) Training-cohort dependence of off-diagonal transfer.",
        "Figure 6. Robustness sensitivity. (A) Feature retention, FDR < 0.10 association counts, and stable-candidate counts across prevalence and cohort-presence filters. (B) Pseudocount sensitivity for direction concordance and FDR < 0.10 feature counts.",
        "Figure 7. Leakage-aware panel benchmarking. (A) LOSO AUROC distributions after training-only feature selection. (B) Median overlap with full-data reference panels. (C) Random-panel fold context; points represent held-out studies and compare observed training-only stability-panel AUROC with the random-panel median AUROC in the same fold. (D) Paired LOSO AUROC differences with bootstrap 95% confidence intervals. These analyses test cohort-shift separability and do not represent clinical validation.",
        "Figure 8. Ecological guild score interpretation. (A) Random-effects Hedges g summaries for taxonomy-defined guild scores. (B) Cohort-level Hedges g spread by guild. (C) Guild membership and overlap with stable candidates. (D) Score-only LOSO separability stress test. Guild panels are taxonomic annotation summaries, not functional or source-attribution models.",
        "Supplementary Figure S1. Covariate completeness and primary sensitivity comparison. (A) Availability of age, BMI, gender, country, platform, and disease-stage fields. (B) Association comparison between baseline and covariate-adjusted models.",
        "Supplementary Figure S2. Stable-candidate evidence map. (A) Leave-one-cohort-out robustness for the 29 stable candidates. (B) Candidate-level support layers: Covariate, Meta-analysis, LOCO, Elastic-net, MicrobiomeHD, and Oral panel.",
    ]
    for legend in legends:
        add_para(doc, legend)

    doc.save(MANUSCRIPT_DOCX)
    write_markdown_from_docx(doc, MANUSCRIPT_MD)


def write_markdown_from_docx(doc: Document, path: Path) -> None:
    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name == "Title":
            lines.append(f"# {text}")
        elif p.style.name == "Heading 1":
            lines.append(f"\n## {text}")
        elif p.style.name == "Heading 2":
            lines.append(f"\n### {text}")
        else:
            lines.append(f"\n{text}")
    path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def build_cover_letter() -> None:
    doc = Document()
    style_document(doc)
    paragraphs = [
        "Dear Editor,",
        f"Please consider my manuscript entitled \"{TITLE}\" for publication as a Research Article in BMC Microbiology.",
        "This revision frames the study not as another pooled CRC microbiome association analysis, but as a transferability-aware candidate prioritization workflow. The manuscript tests whether association-ranked microbial signals, cross-cohort stable signals, and portability-aware panels behave similarly under cohort shift. The addition of Aitchison-space variance partitioning, training-only panel benchmarking, random-panel comparisons, and transparent ecological guild scores strengthens the evidence that pooled association strength and cross-cohort portability should be evaluated separately.",
        "The manuscript is aligned with BMC Microbiology because it addresses human gut microbiome signals using a transparent and reproducible computational workflow on sample-level public processed metagenome matrices. The study is not a literature-level pooled analysis, bibliometric analysis, raw-read reprocessing study, or clinical diagnostic validation study. It provides bounded processed-data association, robustness, heterogeneity, and transferability evidence for candidate prioritization.",
        "The manuscript is original, is not under consideration elsewhere, and has been approved by the sole author. I declare no competing interests and received no specific funding for this work.",
        "Sincerely,",
        AUTHOR,
        AFFILIATION,
        EMAIL,
    ]
    for paragraph in paragraphs:
        add_para(doc, paragraph)
    doc.save(COVER_DOCX)
    COVER_MD.write_text("\n\n".join(paragraphs) + "\n", encoding="utf-8")


def main() -> None:
    build_manuscript()
    build_cover_letter()
    print({"manuscript_docx": str(MANUSCRIPT_DOCX), "cover_letter_docx": str(COVER_DOCX), "title": TITLE})


if __name__ == "__main__":
    main()

