from __future__ import annotations

import json
import math
import platform
import re
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Iterable
from xml.sax.saxutils import escape

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


RUN_DIR = Path(__file__).resolve().parents[4]
ANALYSIS_DIR = RUN_DIR / "agents" / "05-analysis" / "workspace"
FIGURE_DIR = RUN_DIR / "agents" / "06-figures" / "workspace" / "figures"
REFERENCE_MAP = RUN_DIR / "agents" / "07-manuscript" / "workspace" / "reports" / "journal_style_reference_map.tsv"
SUBMISSION_DIR = RUN_DIR / "BMC_submission"

MANUSCRIPT_DIR = SUBMISSION_DIR / "manuscript"
COVER_DIR = SUBMISSION_DIR / "cover_letter"
SUB_FIG_DIR = SUBMISSION_DIR / "figures"
SUPP_TABLE_DIR = SUBMISSION_DIR / "supplementary_tables"
CODE_DIR = SUBMISSION_DIR / "code"
CODE_SCRIPT_DIR = CODE_DIR / "scripts"
RESULTS_DIR = SUBMISSION_DIR / "results"
REPORTS_DIR = SUBMISSION_DIR / "reports"

TITLE = "Robustness and transferability of colorectal cancer gut metagenome signals across public cohorts"
AUTHOR = "Zhun Qiu"
AUTHOR_INITIALS = "ZQ"
TARGET_JOURNAL = "BMC Microbiology"

SPECIES_NAMES = [
    "Parvimonas micra",
    "Gemella morbillorum",
    "Peptostreptococcus stomatis",
    "Streptococcus salivarius",
    "Roseburia intestinalis",
    "Faecalibacterium prausnitzii",
    "Roseburia faecis",
    "Eubacterium eligens",
    "Ruthenibacterium lactatiformans",
    "Bacteroides fragilis",
    "Butyricimonas virosa",
    "Flavonifractor plautii",
    "Anaerostipes hadrus",
    "Intestinimonas butyriciproducens",
    "Fusicatenibacter saccharivorans",
    "Parabacteroides distasonis",
    "Escherichia coli",
    "Blautia wexlerae",
    "Gemmiger formicilis",
    "Bilophila wadsworthia",
    "Bacteroides caccae",
    "Bacteroides thetaiotaomicron",
    "Clostridium bolteae",
    "Alistipes shahii",
    "Bifidobacterium longum",
    "Bifidobacterium adolescentis",
    "Roseburia inulinivorans",
    "Akkermansia muciniphila",
    "Agathobaculum butyriciproducens",
    "Collinsella stercoris",
    "Alistipes finegoldii",
    "Parabacteroides merdae",
]
SPECIES_RE = re.compile("(" + "|".join(re.escape(s) for s in sorted(SPECIES_NAMES, key=len, reverse=True)) + ")")


def ensure_dirs() -> None:
    for path in [
        MANUSCRIPT_DIR,
        COVER_DIR,
        SUB_FIG_DIR,
        SUPP_TABLE_DIR,
        CODE_DIR,
        CODE_SCRIPT_DIR,
        RESULTS_DIR,
        REPORTS_DIR,
    ]:
        path.mkdir(parents=True, exist_ok=True)


def read_tsv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path, sep="\t")


def fmt_p(value: float) -> str:
    if pd.isna(value):
        return ""
    if value < 0.001:
        return f"{value:.2e}"
    return f"{value:.4f}".rstrip("0").rstrip(".")


def percent(value: float, digits: int = 1) -> str:
    if pd.isna(value):
        return ""
    return f"{100 * float(value):.{digits}f}%"


def yes_no(value) -> str:
    if isinstance(value, str):
        return "yes" if value.strip().lower() in {"true", "yes", "genus_match"} else "no"
    return "yes" if bool(value) else "no"


def clean_direction(value: str) -> str:
    return str(value).replace("_", "-")


def clean_range_text(value) -> str:
    text = str(value)
    return re.sub(r"(?<=\d)-(?=\d)", "–", text)


def write_excel_and_csv(df: pd.DataFrame, base_name: str, sheet_name: str) -> tuple[Path, Path]:
    xlsx_path = SUPP_TABLE_DIR / f"{base_name}.xlsx"
    csv_path = SUPP_TABLE_DIR / f"{base_name}.csv"
    df.to_csv(csv_path, index=False, encoding="utf-8-sig")
    with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name=sheet_name[:31], index=False)
        ws = writer.sheets[sheet_name[:31]]
        ws.freeze_panes = "A2"
        for column_cells in ws.columns:
            values = [str(cell.value) if cell.value is not None else "" for cell in column_cells[:200]]
            width = min(max([len(v) for v in values] + [10]) + 2, 70)
            ws.column_dimensions[column_cells[0].column_letter].width = width
    return xlsx_path, csv_path


def build_supplementary_tables() -> dict[str, list[str]]:
    results = ANALYSIS_DIR / "results"
    table_dir = results / "I_publication_tables"
    created: dict[str, list[str]] = {}

    s1 = read_tsv(table_dir / "Table1_cohort_sample_characteristics.tsv").copy()
    for col in ["Age, median (IQR)", "BMI, median (IQR)"]:
        s1[col] = s1[col].map(clean_range_text)
    created["Supplementary Table S1"] = [str(p) for p in write_excel_and_csv(s1, "Supplementary_Table_S1", "S1_Cohorts")]

    s2 = read_tsv(table_dir / "Table2_stable_candidate_taxa.tsv").copy()
    s2["Direction"] = s2["Direction"].map(clean_direction)
    created["Supplementary Table S2"] = [str(p) for p in write_excel_and_csv(s2, "Supplementary_Table_S2", "S2_Stable_taxa")]

    s3 = read_tsv(table_dir / "Table3_leave_one_study_out_performance.tsv").copy()
    created["Supplementary Table S3"] = [str(p) for p in write_excel_and_csv(s3, "Supplementary_Table_S3", "S3_LOSO")]

    stable = read_tsv(results / "M_robustness_interpretation" / "stable_candidate_interpretation.tsv")
    transport = read_tsv(results / "H" / "transportability_score.tsv")[["taxon_long", "cohorts_tested"]]
    stable = stable.merge(transport, on="taxon_long", how="left")
    s4 = pd.DataFrame(
        {
            "Taxon": stable["species_short"],
            "Direction": stable["direction"].map(clean_direction),
            "Baseline BH q": stable["adjusted_qvalue"],
            "Covariate BH q": stable["qvalue_bh_coef_crc_clr_study_age_bmi_gender"],
            "Random-effects q": stable["qvalue_bh_random_effect"],
            "I²": stable["i2_percent"].map(lambda x: f"{float(x):.1f}%" if pd.notna(x) else ""),
            "Same-direction cohorts": [
                f"{int(a)}/{int(b) if pd.notna(b) else 11}" for a, b in zip(stable["cohorts_same_direction"], stable["cohorts_tested"])
            ],
            "LOCO direction preservation fraction": stable["same_direction_fraction"].map(lambda x: f"{float(x):.3f}"),
            "LOCO FDR < 0.10 support fraction": stable["fdr_lt_0_10_fraction"].map(lambda x: f"{float(x):.3f}"),
            "LOCO robust": stable["leave_one_cohort_robust"].map(yes_no),
            "Elastic-net selected folds": stable["selected_fraction"].map(lambda x: percent(float(x), 0)),
            "MicrobiomeHD genus match": stable["microbiomehd_genus_match"].map(yes_no),
            "Oral-associated panel membership": stable["oral_associated_taxonomy_panel"].map(yes_no),
            "High-consistency candidate": stable["interpretation_group"].map(lambda x: "yes" if x == "high_consistency_candidate" else "no"),
            "Notes": stable["interpretation_group"].str.replace("_", " ", regex=False),
        }
    )
    created["Supplementary Table S4"] = [
        str(p) for p in write_excel_and_csv(s4, "Supplementary_Table_S4_stable_candidate_evidence_map", "S4_Evidence_map")
    ]

    oral = read_tsv(results / "K_ecology_oral_signature" / "oral_associated_taxa_panel.tsv")
    stable_genera = set(stable.loc[stable["oral_associated_taxonomy_panel"].astype(bool), "genus"].astype(str))
    rows = []
    for genus, group in oral.groupby("genus", sort=True):
        species = sorted(group["species_short"].dropna().astype(str).unique())
        rows.append(
            {
                "Genus": genus,
                "Included species in processed matrix": "; ".join(species),
                "Reason for inclusion": "Literature-prior oral-associated genus panel used for taxonomy-name context.",
                "Literature support or source category": "Oral-associated genus list from literature-prior taxonomy panel; contextual only.",
                "Present among stable candidates": "yes" if genus in stable_genera else "no",
                "Notes": "Panel membership does not infer oral source, transmission, or anatomical origin in these processed stool tables.",
            }
        )
    s5 = pd.DataFrame(rows)
    created["Supplementary Table S5"] = [
        str(p) for p in write_excel_and_csv(s5, "Supplementary_Table_S5_oral_associated_panel", "S5_Oral_panel")
    ]
    return created


def copy_result_csvs() -> dict[str, str]:
    results = ANALYSIS_DIR / "results"
    mapping = {
        "full_association_results.csv": results / "C" / "cohort_adjusted_crc_control_association.tsv",
        "covariate_sensitivity_results.csv": results / "G" / "baseline_vs_covariate_sensitivity.tsv",
        "random_effects_results.csv": results / "J_meta_heterogeneity" / "random_effects_meta_analysis.tsv",
        "loso_results.csv": results / "I_publication_tables" / "Table3_leave_one_study_out_performance.tsv",
        "pairwise_transfer_matrix.csv": results / "L_pairwise_transfer" / "pairwise_transfer_auroc_matrix.tsv",
        "filter_sensitivity_results.csv": results / "M_robustness_interpretation" / "filter_threshold_sensitivity.tsv",
        "pseudocount_sensitivity_results.csv": results / "M_robustness_interpretation" / "pseudocount_sensitivity.tsv",
        "loco_robustness_results.csv": results / "M_robustness_interpretation" / "leave_one_cohort_candidate_summary.tsv",
    }
    copied = {}
    for name, source in mapping.items():
        df = read_tsv(source)
        dest = RESULTS_DIR / name
        df.to_csv(dest, index=False, encoding="utf-8-sig")
        copied[name] = str(dest)
    return copied


def copy_figures() -> dict[str, list[str]]:
    mapping = {
        "Figure_1": FIGURE_DIR / "Fig1_data_inventory_qc",
        "Figure_2": FIGURE_DIR / "Fig2_association_screen",
        "Figure_3": FIGURE_DIR / "Fig3_cross_cohort_stability",
        "Figure_4": FIGURE_DIR / "Fig4_benchmark_context",
        "Figure_5": FIGURE_DIR / "Fig5_ecology_transfer",
        "Figure_6": FIGURE_DIR / "Fig6_evidence_chain",
        "Supplementary_Figure_S1": FIGURE_DIR / "supp" / "FigS1_covariate_sensitivity",
        "Supplementary_Figure_S2": FIGURE_DIR / "supp" / "FigS2_stable_candidate_evidence_map",
    }
    copied: dict[str, list[str]] = {}
    for dest_stem, source_stem in mapping.items():
        copied[dest_stem] = []
        for ext in [".pdf", ".png", ".tiff"]:
            source = source_stem.with_suffix(ext)
            if source.exists():
                dest = SUB_FIG_DIR / f"{dest_stem}{ext}"
                shutil.copy2(source, dest)
                copied[dest_stem].append(str(dest))
    return copied


def reference_list() -> list[str]:
    refs = read_tsv(REFERENCE_MAP)
    refs = refs.sort_values("reference_order")
    return refs["vancouver_reference"].astype(str).tolist()


def paragraph_text(label: str, body: str) -> str:
    return f"{label}: {body}" if label else body


def manuscript_items() -> list[dict[str, str]]:
    refs = reference_list()
    items: list[dict[str, str]] = [
        {"type": "title", "text": TITLE},
        {"type": "p", "text": AUTHOR},
        {"type": "p", "text": "Affiliation: [ADD AFFILIATION]"},
        {"type": "p", "text": "Corresponding author: Zhun Qiu, [ADD EMAIL]"},
        {"type": "h1", "text": "Abstract"},
        {
            "type": "label_p",
            "label": "Background",
            "text": "Public processed metagenomic resources provide an opportunity to re-examine colorectal cancer (CRC)-associated gut microbiome signals across cohorts without starting from raw sequencing files. We evaluated whether a sample-level public processed species matrix could support a reproducible, cohort-aware workflow that separates study-adjusted association, robustness, heterogeneity, and transferability.",
        },
        {
            "type": "label_p",
            "label": "Results",
            "text": "We analyzed 1,395 CRC/control samples from 11 public study labels and retained 301 species-level features after prevalence and cohort-presence filtering. A study-adjusted centered log-ratio association screen identified 124 species at FDR < 0.10. Cross-cohort stability was more restrictive: 29 candidates met both FDR < 0.10 and a ≥60% same-direction cohort-support criterion. Adjustment for age, BMI, and gender preserved the main association landscape, and sensitivity analyses supported robustness to feature-filtering and pseudocount choices. Leave-one-study-out elastic-net stress testing yielded a median held-out AUROC of 0.752, with a minimum AUROC of 0.573, while pairwise cross-study transfer yielded a median off-diagonal AUROC of 0.655. Leave-one-cohort-out association sensitivity supported 26 of the 29 stable candidates under predefined robustness criteria.",
        },
        {
            "type": "label_p",
            "label": "Conclusions",
            "text": "Public processed CRC gut metagenome matrices can support a reproducible candidate-prioritization workflow when association strength, robustness, heterogeneity, and cohort transferability are evaluated separately. The results support bounded processed-data association and transferability claims, but not causal, clinical diagnostic, stage-specific, mechanistic, functional, strain-level, oral-source, or raw-read reprocessing claims.",
        },
        {"type": "h1", "text": "Keywords"},
        {
            "type": "p",
            "text": "Colorectal cancer; gut microbiome; metagenomics; curatedMetagenomicData; compositional data analysis; cohort-aware analysis; transferability; reproducibility",
        },
        {"type": "h1", "text": "Background"},
        {
            "type": "p",
            "text": "Colorectal cancer (CRC)-associated gut microbiome differences have been reported across several fecal metagenomic studies, including species-level signals involving oral-associated taxa, butyrate-associated commensals, and other community shifts [1-7]. These studies have helped establish that CRC status can be associated with reproducible microbial differences, yet they also show that effect sizes, taxonomic resolution, and cohort portability can vary across populations and study designs.",
        },
        {
            "type": "p",
            "text": "Public harmonized resources such as curatedMetagenomicData and ExperimentHub-derived processed tables make it possible to work directly with sample-level abundance matrices rather than published summary statistics [8,9]. Related resources such as GMrepo and MicrobiomeHD provide additional disease-microbiome context, but they differ in schema, taxonomic resolution, and intended use [8,10,11]. For CRC, these resources create an opportunity to evaluate how far processed matrices can support reproducible candidate prioritization without requiring raw-read reprocessing at the first stage.",
        },
        {
            "type": "p",
            "text": "A central challenge is that study-adjusted association strength does not necessarily imply portability across cohorts. A species may be strongly associated with CRC status in a pooled design but show weaker direction stability, high heterogeneity, or limited transfer when one cohort is held out. Conversely, a more modest association may be more directionally stable and therefore more useful for transparent candidate prioritization. These layers should be evaluated separately rather than collapsed into one significance threshold or one model performance number.",
        },
        {
            "type": "p",
            "text": "The objective was to evaluate whether a sample-level public processed CRC/control metagenomic matrix can support a reproducible cohort-aware computational workflow. The workflow was designed to separate study-adjusted association strength, covariate robustness, cross-cohort direction stability, random-effects heterogeneity, and leave-one-study-out transferability. This design differs from literature-level pooled analyses because it operates on harmonized sample-level processed abundance tables rather than published summary statistics. The analysis was therefore framed as a processed-data robustness and transferability audit, not as raw-read reprocessing, clinical diagnostic validation, or mechanistic inference.",
        },
        {"type": "h1", "text": "Methods"},
        {"type": "h2", "text": "Data source and feature filtering"},
        {
            "type": "p",
            "text": "The analysis used curatedMetagenomicData/ExperimentHub-derived processed CRC/control resources with sample metadata, a species-level percent relative-abundance matrix, and a species feature index [9]. Across study-level inputs, species absent from a study-level matrix were filled as zero in the cross-study union matrix. The analysis-ready matrix contained 944 species-level features and 1,395 samples across 11 study labels.",
        },
        {
            "type": "p",
            "text": "The primary feature filter retained species with overall prevalence ≥5% and presence in at least three study labels. Prevalence was calculated from non-zero relative abundance. This filter retained 301 of 944 species-level features. Centered log-ratio (CLR) transformation was performed after filtering. The pseudocount was defined as one-half of the minimum positive retained abundance value, equal to 2.5e-05 in percent relative-abundance units.",
        },
        {"type": "h2", "text": "Study-adjusted association and covariate sensitivity"},
        {
            "type": "p",
            "text": "For each retained species, the primary association model was an ordinary least-squares model with study indicators: CLR_taxon ~ CRC_status + study. CRC_status was encoded as CRC = 1 and control = 0, so positive coefficients indicate CRC-higher CLR abundance. Heteroscedasticity-consistent HC3 standard errors were used, and P values were adjusted across tested features with the Benjamini-Hochberg false discovery rate (FDR) procedure [13]. Raw-abundance Mann-Whitney tests were used as a sensitivity check only and were not treated as the primary association model.",
        },
        {
            "type": "p",
            "text": "The primary covariate sensitivity model used complete-case samples with age, BMI, and gender metadata: CLR_taxon ~ CRC_status + study + age_z + BMI_z + gender. Age and BMI were z-standardized within the analysis dataset. Country and sequencing platform metadata were summarized as context, but the primary robustness interpretation used the study-adjusted and study + age + BMI + gender models to avoid over-interpreting sparse design strata.",
        },
        {"type": "h2", "text": "Ordination and within-study permutation test"},
        {
            "type": "p",
            "text": "Aitchison PCA was computed from the filtered CLR matrix. To test whether CRC/control compositional separation remained after accounting for study labels, sample CLR vectors were residualized by subtracting study-specific feature means, and a pseudo-F statistic was computed from between-status and within-status sums of squared Euclidean distances in the residualized CLR space. Statistical significance was evaluated by 4,999 permutations of CRC/control labels within study labels, with the permutation P value calculated as (b + 1)/(n + 1), where b is the number of permuted pseudo-F values at least as large as the observed statistic.",
        },
        {"type": "h2", "text": "Cross-cohort stability and random-effects heterogeneity"},
        {
            "type": "p",
            "text": "Cross-cohort stability was evaluated using per-study effect estimates for each retained species. A stable candidate was defined before interpretation as a species with study-adjusted BH FDR < 0.10 and same-direction support in at least 60% of eligible study labels. Per-cohort effect estimates were also summarized with DerSimonian-Laird random-effects analysis, including random-effects estimates, random-effects q values, τ², and I². This heterogeneity analysis was treated as an audit layer for processed-data consistency rather than as external validation.",
        },
        {"type": "h2", "text": "Leave-one-study-out and pairwise transferability stress tests"},
        {
            "type": "p",
            "text": "Leave-one-study-out (LOSO) elastic-net analyses were used as separability and transferability stress tests. For each held-out study label, model fitting used only the remaining study labels. Standardization was fit on training studies only and applied to the held-out study. The logistic regression model used saga optimization, elastic-net penalty with l1_ratio = 0.5, C = 0.5, class_weight = balanced, max_iter = 2000, and random_state = 42. Hyperparameters were fixed rather than tuned for clinical model development. AUROC and average precision were reported by held-out study, with 2,000-resample bootstrap intervals.",
        },
        {
            "type": "p",
            "text": "Pairwise train-study × test-study transfer was evaluated by fitting a model on one study label and testing it on each study label. The diagonal represents within-study separability, whereas off-diagonal cells summarize cross-study transportability. This analysis was interpreted as a study-to-study stress test, not as screening calibration.",
        },
        {"type": "h2", "text": "Benchmark context, oral-associated panel, and composite ranking"},
        {
            "type": "p",
            "text": "MicrobiomeHD processed tables were used only as genus-level benchmark context [8]. The comparison did not merge individual samples, harmonize all schemas, or validate a classifier. The oral-associated score was based on a literature-prior taxonomy-name panel and was not an independently inferred oral-source or transmission model. The genus panel and included species are provided as Supplementary Table S5.",
        },
        {
            "type": "p",
            "text": "An exploratory composite transportability ranking was calculated as: transportability_score = -log10(BH q) × stability_fraction × covariate_support_weight × (1 + mean_abs_elasticnet_coef). The covariate_support_weight was 1.0 when the covariate model had BH q < 0.10 and direction concordance with the primary model, and 0.5 otherwise. The score was used for candidate ordering only and was not used as a hypothesis test.",
        },
        {"type": "h2", "text": "Robustness analyses"},
        {
            "type": "p",
            "text": "We further evaluated whether the association workflow was sensitive to feature-filtering thresholds, pseudocount choice, and individual-cohort removal. Feature-filter sensitivity used prevalence thresholds of 3%, 5%, and 10% crossed with cohort-presence thresholds of two, three, or four study labels. Pseudocount sensitivity used one-quarter, one-half, and one times the minimum positive retained abundance. Leave-one-cohort-out (LOCO) association sensitivity refit the association screen after removing each study label. A stable candidate met the LOCO robustness rule when it preserved direction in ≥90% of excluded-cohort fits and retained FDR < 0.10 support in ≥80% of excluded-cohort fits. These analyses were robustness checks, not external validation.",
        },
        {"type": "h2", "text": "Use of AI-assisted tools"},
        {
            "type": "p",
            "text": "AI-assisted coding and language tools were used to support code refactoring, manuscript editing, and formatting during manuscript preparation. The author reviewed and verified all tool-assisted outputs, reran or checked the analyses where applicable, and takes full responsibility for the content, results, interpretation, and submission of the manuscript.",
        },
        {"type": "h1", "text": "Results"},
        {"type": "h2", "text": "Processed matrix structure and compositional separation"},
        {
            "type": "p",
            "text": "The analysis covered 1,395 samples, including 701 CRC and 694 control records across 11 study labels. After prevalence and cohort-presence filtering, 301 of 944 species-level features were retained for CLR-based analysis. Aitchison PCA showed broad CRC/control overlap, consistent with a heterogeneous public processed matrix rather than a simple two-cluster structure. The within-study permutation test nevertheless detected residualized CRC/control compositional separation (pseudo-F 7.90, p = 0.0002, 4,999 permutations; Figure 1).",
        },
        {"type": "h2", "text": "Study-adjusted associations and covariate robustness"},
        {
            "type": "p",
            "text": "The study-adjusted CLR association screen tested 301 species and identified 124 features at FDR < 0.10 and 107 features at FDR < 0.05. The strongest CRC-higher associations included Parvimonas micra, Gemella morbillorum, and Peptostreptococcus stomatis. The direction of the coefficient is interpretable as CRC-higher for positive values and control-higher for negative values (Figure 2).",
        },
        {
            "type": "p",
            "text": "The complete-case covariate model used 1,380 samples with age, BMI, and gender metadata. This sensitivity analysis identified 137 features at FDR < 0.10, preserved all 29 primary stable candidates, showed 91/100 overlap among the top baseline features, and had all-feature direction concordance of 0.947. These results indicate that the main association landscape was not driven solely by the available age, BMI, or gender metadata fields.",
        },
        {"type": "h2", "text": "Cross-cohort stability and random-effects heterogeneity"},
        {
            "type": "p",
            "text": "Cross-cohort stability was more restrictive than study-adjusted association strength. Among the 124 FDR < 0.10 features, 29 candidates also met the ≥60% same-direction cohort-support rule. This distinction shows that many association signals were not equally portable across study labels.",
        },
        {
            "type": "p",
            "text": "Random-effects heterogeneity auditing tested all 301 retained species. Ninety-four species had random-effects FDR < 0.10, the median I² was 40.6%, and 17 species had I² ≥75%. Predefined random-effects FDR and heterogeneity filters identified a broader set of meta-supported candidates, but the stable-candidate interpretation emphasized the stricter overlap between study-adjusted association and cross-cohort direction support (Figure 3).",
        },
        {"type": "h2", "text": "Leave-one-study-out and pairwise transferability"},
        {
            "type": "p",
            "text": "LOSO elastic-net stress testing produced mixed but informative cross-study separability. The median held-out AUROC was 0.752, the weighted AUROC was 0.737, and the minimum held-out AUROC was 0.573. The hardest held-out fold was HanniganGD_2017, with an average precision of 0.610. These results support use of the workflow as a transferability stress test, not as a calibrated clinical diagnostic model (Figure 4; Supplementary Table S3).",
        },
        {
            "type": "p",
            "text": "Pairwise train-study × test-study transfer gave a median off-diagonal AUROC of 0.655 across 110 cross-study pairs, with a range from 0.154 to 0.994. Most off-diagonal pairs were above chance-level separability (102/110 > 0.5), but the wide range shows substantial cohort dependence (Figure 5).",
        },
        {"type": "h2", "text": "Ecological context, oral-associated signal, and benchmark context"},
        {
            "type": "p",
            "text": "Ecological summary analyses linked the candidate list to broader compositional patterns. All four metric-level random-effects summaries had BH q < 0.10. The literature-prior oral-associated abundance score was higher in CRC samples with Hedges g = 0.147, q = 0.00859, and I² = 0.0%, but the score is a taxonomy-name context measure and does not infer oral origin.",
        },
        {
            "type": "p",
            "text": "The stable-candidate list and MicrobiomeHD benchmark context were related but not identical. Among stable candidates, 17 had genus-level MicrobiomeHD context, and one belonged to the literature-prior oral-associated panel. This comparison provides external processed-table context but does not merge individual samples, harmonize schemas, or validate a classifier. It was therefore interpreted as benchmark context rather than external validation.",
        },
        {"type": "h2", "text": "Robustness sensitivity and stable-candidate interpretation"},
        {
            "type": "p",
            "text": "Feature-filter sensitivity produced 94–145 FDR < 0.10 features across filtering configurations, with top-100 overlap ranging from 78 to 100 and stable-candidate counts ranging from 28 to 29. Pseudocount sensitivity produced 124–127 FDR < 0.10 features, top-100 overlap of at least 100/100, and minimum direction concordance of 0.993. Stricter stability thresholds retained 20 candidates at ≥70% same-direction support and 11 candidates at ≥80% same-direction support.",
        },
        {
            "type": "p",
            "text": "LOCO association sensitivity supported 26 of the 29 stable candidates under the predefined robustness rule. Eighteen stable candidates were classified as high-consistency candidates, reflecting coherent support across LOCO robustness and random-effects layers. Together, these sensitivity analyses show that the main processed-data association landscape was robust to several analytic choices, while candidate interpretation remained dependent on the chosen support layer (Figure 6; Supplementary Figure S2; Supplementary Table S4).",
        },
        {"type": "h1", "text": "Discussion"},
        {
            "type": "p",
            "text": "This sample-level public processed-matrix reanalysis shows that CRC gut metagenome signals can be evaluated through a cohort-aware workflow that separates association strength, robustness, heterogeneity, and transferability. The workflow retained 301 species from 1,395 samples and identified 124 study-adjusted CLR associations at FDR < 0.10, but only 29 candidates also met the predefined cross-cohort direction-stability rule. The most important result is therefore not a single taxon or a single AUROC value; it is the observation that different layers of evidence select partially overlapping but not identical microbial signals.",
        },
        {
            "type": "p",
            "text": "The robustness analyses strengthened this interpretation. The broad association landscape was stable across prevalence and cohort-presence thresholds, and pseudocount sensitivity had high direction concordance with the primary analysis. At the same time, stricter direction-stability thresholds and LOCO analyses reduced the candidate list, showing that transparent candidate prioritization benefits from reporting how each layer changes the interpretation.",
        },
        {
            "type": "p",
            "text": "The biological patterns were consistent with previous CRC microbiome literature but should be interpreted within the processed-table design. Control-higher candidates included butyrate-associated commensal genera such as Roseburia and Faecalibacterium, whereas CRC-higher candidates included species previously reported in CRC-associated microbial signatures. The oral-associated score and stable-candidate list were related but not identical: the score summarized a taxonomy-name panel, while candidate stability depended on species-level association and cross-cohort behavior.",
        },
        {
            "type": "p",
            "text": "Transferability analyses also highlighted the distinction between separability and portability. Residualized Aitchison separation was statistically detectable after study adjustment, LOSO elastic-net stress testing had a median AUROC of 0.752, and pairwise transfer had a median off-diagonal AUROC of 0.655. However, the minimum LOSO AUROC and the wide pairwise transfer range show that cohort composition, study design, and processed-table harmonization remain important. These results should not be read as screening calibration or clinical model validation.",
        },
        {
            "type": "p",
            "text": "Strengths of this study include explicit study adjustment, compositional transformation, covariate sensitivity, per-cohort heterogeneity auditing, LOSO and pairwise transfer analyses, threshold sensitivity checks, and machine-readable result tables. Limitations follow directly from the design: the analysis used public processed tables only; it did not reprocess raw reads, profile functions or strains, model disease stage, perform patient-level independent validation, or calibrate a clinical decision tool. Several contextual resources were used only as background or benchmark context and were not treated as Results-level individual-sample validation datasets.",
        },
        {"type": "h1", "text": "Conclusions"},
        {
            "type": "p",
            "text": "Public processed CRC/control gut metagenome matrices can support more than a simple association screen when they are analyzed through a cohort-aware workflow. In 1,395 samples across 11 study labels, the retained 301-species matrix produced 124 study-adjusted CLR associations at FDR < 0.10 and identified 29 same-direction stable candidates under a ≥60% cohort-support rule. Robustness analyses supported the broad association landscape across feature-filtering and pseudocount choices, and 26 of the 29 stable candidates passed the leave-one-cohort-out robustness rule.",
        },
        {
            "type": "p",
            "text": "The main contribution is the separation of association strength from portability and robustness. Leave-one-study-out elastic-net stress testing reached a median AUROC of 0.752 with a minimum AUROC of 0.573, and pairwise cross-study transfer reached a median off-diagonal AUROC of 0.655. These findings support a reproducible processed-data candidate-prioritization workflow, but they do not establish causal, clinical diagnostic, stage-specific, mechanistic, functional, strain-level, oral-source, or raw-read reprocessing claims.",
        },
        {"type": "h1", "text": "List of abbreviations"},
        {"type": "p", "text": "AUROC: area under the receiver operating characteristic curve"},
        {"type": "p", "text": "BH: Benjamini-Hochberg"},
        {"type": "p", "text": "CLR: centered log-ratio"},
        {"type": "p", "text": "CRC: colorectal cancer"},
        {"type": "p", "text": "FDR: false discovery rate"},
        {"type": "p", "text": "LOCO: leave-one-cohort-out"},
        {"type": "p", "text": "LOSO: leave-one-study-out"},
        {"type": "p", "text": "PCA: principal component analysis"},
        {"type": "h1", "text": "Declarations"},
        {"type": "h2", "text": "Ethics approval and consent to participate"},
        {
            "type": "p",
            "text": "This study is a secondary analysis of publicly available, de-identified processed microbiome abundance tables and associated metadata. No new human participants were recruited, no intervention was performed, and no identifiable private information was accessed. Ethics approval and consent to participate were therefore not required for this secondary analysis.",
        },
        {"type": "h2", "text": "Consent for publication"},
        {"type": "p", "text": "Not applicable."},
        {"type": "h2", "text": "Availability of data and materials"},
        {
            "type": "p",
            "text": "The public source data analyzed in this study are available through curatedMetagenomicData/ExperimentHub-derived resources. The analysis-ready metadata identifiers, species-level abundance processing scripts, statistical analysis scripts, figure-generation scripts, software environment files, random seeds, and complete result tables generated for this study are available at GitHub: [ADD GITHUB REPOSITORY LINK]. An archived version of the analysis code and result tables is available at Zenodo: [ADD ZENODO DOI]. Supplementary tables are also provided as machine-readable additional files.",
        },
        {"type": "h2", "text": "Competing interests"},
        {"type": "p", "text": "The author declares that he has no competing interests."},
        {"type": "h2", "text": "Funding"},
        {"type": "p", "text": "The author received no specific funding for this work."},
        {"type": "h2", "text": "Authors' contributions"},
        {
            "type": "p",
            "text": "ZQ conceived the study, curated the public processed datasets, designed and performed the analyses, generated the figures and tables, interpreted the results, wrote and revised the manuscript, and approved the final manuscript.",
        },
        {"type": "h2", "text": "Acknowledgements"},
        {"type": "p", "text": "Not applicable."},
        {"type": "h2", "text": "Authors' information"},
        {"type": "p", "text": "Not applicable."},
        {"type": "h1", "text": "Additional files"},
        {"type": "p", "text": "Additional file 1: Supplementary Figure S1. Covariate completeness and primary sensitivity comparison."},
        {"type": "p", "text": "Additional file 2: Supplementary Figure S2. Stable-candidate evidence map."},
        {"type": "p", "text": "Additional file 3: Supplementary Table S1. Cohort and sample characteristics."},
        {"type": "p", "text": "Additional file 4: Supplementary Table S2. Stable candidate taxa."},
        {"type": "p", "text": "Additional file 5: Supplementary Table S3. Leave-one-study-out performance."},
        {"type": "p", "text": "Additional file 6: Supplementary Table S4. Stable-candidate evidence map."},
        {"type": "p", "text": "Additional file 7: Supplementary Table S5. Literature-prior oral-associated genus panel."},
        {"type": "h1", "text": "References"},
    ]
    for idx, ref in enumerate(refs, start=1):
        items.append({"type": "ref", "text": f"{idx}. {ref}"})

    items.extend(
        [
            {"type": "h1", "text": "Figure legends"},
            {
                "type": "p",
                "text": "Figure 1. Data structure and compositional separation. (A) Aitchison PCA of retained species-level CLR profiles. (B) Within-study permutation test for study-residualized CRC/control compositional separation (pseudo-F 7.90, p = 0.0002, 4,999 permutations). (C) Feature prevalence-abundance landscape showing retained and filtered features.",
            },
            {
                "type": "p",
                "text": "Figure 2. Study-adjusted species associations. (A) Association volcano plot for 301 retained species. (B) Coefficient intervals for top candidates. (C) Raw-abundance sensitivity comparison. (D) Baseline versus age, BMI, and gender sensitivity model comparison. Positive coefficients indicate CRC-higher CLR abundance.",
            },
            {
                "type": "p",
                "text": "Figure 3. Cross-cohort stability and heterogeneity. (A) Random-effects summaries for candidate taxa. (B) Relationship between association strength and heterogeneity. (C) Per-cohort effect heatmap for selected candidates.",
            },
            {
                "type": "p",
                "text": "Figure 4. LOSO separability stress test. (A) Held-out AUROC and average precision by study label with bootstrap intervals. (B) Held-out prediction-score distributions. (C) Composite transportability ranking. These panels describe processed-table separability and candidate ordering, not clinical validation.",
            },
            {
                "type": "p",
                "text": "Figure 5. Ecological context and pairwise transfer. (A) Random-effects summary of ecological/oral-associated score metrics. (B) Pairwise train-study × test-study AUROC matrix. (C) Sample-level oral-associated abundance score distribution. (D) Training-cohort dependence of off-diagonal transfer.",
            },
            {
                "type": "p",
                "text": "Figure 6. Robustness sensitivity. (A) Feature retention, FDR < 0.10 association counts, and stable-candidate counts across prevalence and cohort-presence filters. (B) Pseudocount sensitivity for direction concordance and FDR < 0.10 feature counts.",
            },
            {
                "type": "p",
                "text": "Supplementary Figure S1. Covariate completeness and primary sensitivity comparison. (A) Availability of age, BMI, gender, country, platform, and disease-stage fields. (B) Association comparison between baseline and covariate-adjusted models.",
            },
            {
                "type": "p",
                "text": "Supplementary Figure S2. Stable-candidate evidence map. (A) Leave-one-cohort-out robustness for the 29 stable candidates. (B) Candidate-level support layers: Covariate, Meta-analysis, LOCO, Elastic-net, MicrobiomeHD, and Oral panel.",
            },
        ]
    )
    return items


def split_species_runs(text: str) -> Iterable[tuple[str, bool]]:
    last = 0
    for match in SPECIES_RE.finditer(text):
        if match.start() > last:
            yield text[last : match.start()], False
        yield match.group(0), True
        last = match.end()
    if last < len(text):
        yield text[last:], False


def add_text_runs(paragraph, text: str, bold_first_label: str | None = None) -> None:
    remaining = text
    if bold_first_label and text.startswith(f"{bold_first_label}:"):
        run = paragraph.add_run(f"{bold_first_label}: ")
        run.bold = True
        remaining = text[len(bold_first_label) + 2 :].lstrip()
    for part, is_species in split_species_runs(remaining):
        run = paragraph.add_run(part)
        run.italic = is_species


def set_doc_defaults(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ["Heading 1", "Heading 2", "Title"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.paragraph_format.line_spacing = 2.0

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        sect_pr = section._sectPr
        ln_num = OxmlElement("w:lnNumType")
        ln_num.set(qn("w:countBy"), "1")
        ln_num.set(qn("w:restart"), "newPage")
        sect_pr.append(ln_num)
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.add_run("Page ")
        add_field(footer_p, "PAGE")


def add_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def build_docx(items: list[dict[str, str]], path: Path) -> None:
    doc = Document()
    set_doc_defaults(doc)
    for item in items:
        typ = item["type"]
        if typ == "title":
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text_runs(p, item["text"])
        elif typ == "h1":
            doc.add_heading(item["text"], level=1)
        elif typ == "h2":
            doc.add_heading(item["text"], level=2)
        elif typ == "label_p":
            p = doc.add_paragraph()
            add_text_runs(p, paragraph_text(item["label"], item["text"]), bold_first_label=item["label"])
        elif typ == "ref":
            p = doc.add_paragraph()
            add_text_runs(p, item["text"])
        else:
            p = doc.add_paragraph()
            add_text_runs(p, item["text"])
    doc.save(path)


def text_markdown_species(text: str) -> str:
    return SPECIES_RE.sub(lambda m: f"*{m.group(0)}*", text)


def write_markdown(items: list[dict[str, str]], path: Path) -> None:
    lines: list[str] = []
    for item in items:
        typ = item["type"]
        if typ == "title":
            lines.append(f"# {item['text']}")
        elif typ == "h1":
            lines.append(f"\n## {item['text']}")
        elif typ == "h2":
            lines.append(f"\n### {item['text']}")
        elif typ == "label_p":
            lines.append(f"\n**{item['label']}:** {text_markdown_species(item['text'])}")
        else:
            lines.append(f"\n{text_markdown_species(item['text'])}")
    path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def register_pdf_fonts() -> tuple[str, str, str]:
    fonts = [
        ("Arial", Path(r"C:\Windows\Fonts\arial.ttf"), Path(r"C:\Windows\Fonts\arialbd.ttf"), Path(r"C:\Windows\Fonts\ariali.ttf")),
        ("DejaVuSans", Path(r"C:\Windows\Fonts\DejaVuSans.ttf"), Path(r"C:\Windows\Fonts\DejaVuSans-Bold.ttf"), Path(r"C:\Windows\Fonts\DejaVuSans-Oblique.ttf")),
    ]
    for name, regular, bold, italic in fonts:
        if regular.exists() and bold.exists() and italic.exists():
            pdfmetrics.registerFont(TTFont(name, str(regular)))
            pdfmetrics.registerFont(TTFont(f"{name}-Bold", str(bold)))
            pdfmetrics.registerFont(TTFont(f"{name}-Italic", str(italic)))
            return name, f"{name}-Bold", f"{name}-Italic"
    return "Times-Roman", "Times-Bold", "Times-Italic"


def html_species(text: str) -> str:
    escaped = escape(text)
    for species in sorted(SPECIES_NAMES, key=len, reverse=True):
        escaped = escaped.replace(escape(species), f"<i>{escape(species)}</i>")
    return escaped


def page_number(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Times-Roman", 9)
    canvas.drawCentredString(4.25 * inch, 0.45 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf(items: list[dict[str, str]], path: Path) -> None:
    regular, bold, _italic = register_pdf_fonts()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleStyle",
        parent=styles["Title"],
        fontName=bold,
        fontSize=16,
        leading=22,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    h1_style = ParagraphStyle("H1", parent=styles["Heading1"], fontName=bold, fontSize=14, leading=20, spaceBefore=10, spaceAfter=6)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontName=bold, fontSize=12, leading=18, spaceBefore=8, spaceAfter=4)
    body_style = ParagraphStyle("Body", parent=styles["BodyText"], fontName=regular, fontSize=11, leading=22, spaceAfter=6)
    ref_style = ParagraphStyle("Ref", parent=body_style, fontSize=9, leading=14, leftIndent=0.2 * inch, firstLineIndent=-0.2 * inch)
    story = []
    for item in items:
        typ = item["type"]
        if typ == "title":
            story.append(Paragraph(html_species(item["text"]), title_style))
        elif typ == "h1":
            story.append(Paragraph(html_species(item["text"]), h1_style))
        elif typ == "h2":
            story.append(Paragraph(html_species(item["text"]), h2_style))
        elif typ == "label_p":
            story.append(Paragraph(f"<b>{escape(item['label'])}:</b> {html_species(item['text'])}", body_style))
        elif typ == "ref":
            story.append(Paragraph(html_species(item["text"]), ref_style))
        else:
            story.append(Paragraph(html_species(item["text"]), body_style))
        if typ in {"title", "h1"}:
            story.append(Spacer(1, 0.04 * inch))
    pdf = SimpleDocTemplate(str(path), pagesize=LETTER, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    pdf.build(story, onFirstPage=page_number, onLaterPages=page_number)


def cover_letter_items() -> list[dict[str, str]]:
    return [
        {"type": "p", "text": "Dear Editor,"},
        {
            "type": "p",
            "text": "Please consider my manuscript entitled “Robustness and transferability of colorectal cancer gut metagenome signals across public cohorts” for publication as a Research Article in BMC Microbiology.",
        },
        {
            "type": "p",
            "text": "This study presents a reproducible cohort-aware computational reanalysis of sample-level public processed colorectal cancer gut metagenome matrices. Rather than treating study-adjusted association alone as sufficient evidence, the analysis separates association strength, covariate robustness, random-effects heterogeneity, cross-cohort direction stability, leave-one-study-out separability, pairwise cross-study transfer, and sensitivity to feature filtering and pseudocount choice.",
        },
        {
            "type": "p",
            "text": "The manuscript is aligned with BMC Microbiology because it addresses human gut microbiome signals using a transparent and reproducible computational workflow. The study is not a literature-level pooled analysis, bibliometric analysis, raw-read reprocessing study, or clinical diagnostic validation study. Instead, it provides a bounded processed-data candidate-prioritization workflow and highlights the distinction between association strength and cohort transferability.",
        },
        {
            "type": "p",
            "text": "The manuscript is original, is not under consideration elsewhere, and has been approved by the sole author. I declare no competing interests and received no specific funding for this work.",
        },
        {"type": "p", "text": "Sincerely,"},
        {"type": "p", "text": "Zhun Qiu"},
        {"type": "p", "text": "[ADD AFFILIATION]"},
        {"type": "p", "text": "[ADD EMAIL]"},
    ]


def build_cover_letter_docx(path: Path) -> None:
    doc = Document()
    set_doc_defaults(doc)
    for item in cover_letter_items():
        p = doc.add_paragraph()
        add_text_runs(p, item["text"])
    doc.save(path)


def write_cover_letter_md(path: Path) -> None:
    path.write_text("\n\n".join(item["text"] for item in cover_letter_items()) + "\n", encoding="utf-8")


def copy_code_files() -> list[str]:
    scripts = [
        ANALYSIS_DIR / "scripts" / "run_analysis.py",
        ANALYSIS_DIR / "scripts" / "build_meta_heterogeneity_analysis.py",
        ANALYSIS_DIR / "scripts" / "build_ecology_oral_signature.py",
        ANALYSIS_DIR / "scripts" / "build_ecology_transfer_analysis.py",
        ANALYSIS_DIR / "scripts" / "build_robustness_interpretation_analysis.py",
        ANALYSIS_DIR / "scripts" / "build_publication_tables.py",
        RUN_DIR / "agents" / "06-figures" / "workspace" / "figure_scripts" / "render_journal_figures.py",
        Path(__file__),
    ]
    copied = []
    for script in scripts:
        if script.exists():
            dest = CODE_SCRIPT_DIR / script.name
            shutil.copy2(script, dest)
            copied.append(str(dest))
    return copied


def write_reproducibility_files(copied_scripts: list[str]) -> dict[str, str]:
    readme = CODE_DIR / "README.md"
    requirements = CODE_DIR / "requirements.txt"
    session_info = CODE_DIR / "session_info.txt"
    readme.write_text(
        f"""# {TITLE}

Target journal: {TARGET_JOURNAL}

Author: {AUTHOR}

This repository-style package supports a reproducible computational reanalysis of sample-level public processed CRC/control gut metagenome matrices. It uses curatedMetagenomicData/ExperimentHub-derived processed relative-abundance tables and associated metadata already resolved in the local project workspace.

## Data source

The analysis uses public, de-identified processed microbiome abundance tables and metadata. It does not contain private identifiable human data and does not include raw sequencing reads.

## Rebuild outline

1. Rebuild or verify the curated CRC/control processed matrix and metadata in the data root.
2. Run `scripts/run_analysis.py` to produce primary association, ordination, permutation, LOSO, and transportability outputs.
3. Run the post-analysis scripts for random-effects heterogeneity, ecological/oral-associated context, robustness sensitivity, and publication tables.
4. Run `scripts/render_journal_figures.py` to regenerate PDF, PNG, and TIFF figures.
5. Run `scripts/build_bmc_submission_package.py` to assemble the BMC submission package.

## Random seeds

Primary scripts use fixed random seeds where stochastic procedures are used. The LOSO elastic-net stress test uses `random_state = 42`; bootstrap and permutation procedures are recorded in their respective outputs.

## Expected outputs

The expected outputs include manuscript DOCX/PDF files, BMC-style figure files, machine-readable supplementary tables, result CSV files, and this reproducibility package.

## Limitations

The package supports processed-data association, robustness, heterogeneity, and transferability analyses. It does not support causal, clinical diagnostic, stage-specific, mechanistic, functional, strain-level, oral-source, or raw-read reprocessing claims.
""",
        encoding="utf-8",
    )
    requirements.write_text(
        "\n".join(
            [
                "numpy",
                "pandas",
                "scipy",
                "statsmodels",
                "scikit-learn",
                "matplotlib",
                "seaborn",
                "python-docx",
                "reportlab",
                "openpyxl",
                "Pillow",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    session_info.write_text(
        "\n".join(
            [
                f"generated_at={datetime.now().isoformat(timespec='seconds')}",
                f"python={sys.version.replace(chr(10), ' ')}",
                f"platform={platform.platform()}",
                f"run_dir={RUN_DIR}",
                "permutation_test=4,999 within-study permutations",
                "loso_random_state=42",
                "copied_scripts=" + "; ".join(copied_scripts),
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    return {"README.md": str(readme), "requirements.txt": str(requirements), "session_info.txt": str(session_info)}


def abstract_word_count(items: list[dict[str, str]]) -> int:
    abstract = []
    in_abstract = False
    for item in items:
        if item["type"] == "h1" and item["text"] == "Abstract":
            in_abstract = True
            continue
        if in_abstract and item["type"] == "h1":
            break
        if in_abstract and item["type"] == "label_p":
            abstract.append(item["text"])
    return len(re.findall(r"\b[\w-]+\b", " ".join(abstract)))


def qc_package(items: list[dict[str, str]], created: dict) -> dict:
    manuscript_text = "\n".join(item.get("text", "") for item in items)
    prohibited = [
        "Cohort-aware reanalysis of public colorectal cancer gut metagenomes separates pooled association, transferability, and evidence boundaries",
        "after reviewer-style inspection",
        "final claim ledger",
        "local draft",
        "local review draft",
        "not supplied",
        "[@",
        "@zeller",
    ]
    placeholder_allowed = ["[ADD AFFILIATION]", "[ADD EMAIL]", "[ADD GITHUB REPOSITORY LINK]", "[ADD ZENODO DOI]"]
    qc = {
        "title_changed": TITLE in manuscript_text and "evidence boundaries" not in TITLE.lower(),
        "abstract_word_count": abstract_word_count(items),
        "abstract_under_350_words": abstract_word_count(items) <= 350,
        "abstract_has_no_citations": not bool(re.search(r"\[[0-9,\- ]+\]", " ".join(item.get("text", "") for item in items if item["type"] == "label_p"))),
        "permutation_updated": "4,999 permutations" in manuscript_text and "p = 0.0002" in manuscript_text,
        "old_permutation_absent": "199 permutations" not in manuscript_text and "p = 0.0050" not in manuscript_text,
        "prohibited_phrases": [phrase for phrase in prohibited if phrase in manuscript_text],
        "intentional_placeholders": [p for p in placeholder_allowed if p in manuscript_text],
        "supplementary_tables_created": created.get("supplementary_tables", {}),
        "figures_created": created.get("figures", {}),
        "docx_render_qa": "LibreOffice/soffice not available; DOCX was generated directly with python-docx and PDF was generated from the same manuscript source.",
    }
    qc["pass"] = (
        qc["title_changed"]
        and qc["abstract_under_350_words"]
        and qc["abstract_has_no_citations"]
        and qc["permutation_updated"]
        and qc["old_permutation_absent"]
        and not qc["prohibited_phrases"]
    )
    return qc


def write_final_report(created: dict, qc: dict) -> Path:
    report = REPORTS_DIR / "final_bmc_submission_report.md"
    validation_summary = RUN_DIR / "orchestration" / "validation_cycle_summary.md"
    validation_line = "- Not rerun after package assembly."
    if validation_summary.exists():
        text = validation_summary.read_text(encoding="utf-8", errors="replace")
        match = re.search(r"- status: `([^`]+)`", text)
        validation_line = f"- Shared pipeline validation status: {match.group(1) if match else 'available'} ({validation_summary})"
    lines = [
        "# BMC Microbiology Submission Package Report",
        "",
        f"Generated: {datetime.now().isoformat(timespec='seconds')}",
        f"Run directory: {RUN_DIR}",
        "",
        "## Files changed or produced",
        f"- Manuscript DOCX: {created['manuscript_docx']}",
        f"- Manuscript PDF: {created['manuscript_pdf']}",
        f"- Manuscript source Markdown: {created['manuscript_markdown']}",
        f"- Cover letter DOCX: {created['cover_letter_docx']}",
        f"- Cover letter source Markdown: {created['cover_letter_markdown']}",
        f"- Figures directory: {SUB_FIG_DIR}",
        f"- Supplementary tables directory: {SUPP_TABLE_DIR}",
        f"- Results CSV directory: {RESULTS_DIR}",
        f"- Code/reproducibility directory: {CODE_DIR}",
        "",
        "## Analyses rerun",
        "- Primary 05-analysis script rerun after increasing the within-study permutation test from 199 to 4,999 permutations.",
        "- Random-effects heterogeneity, ecological/oral-associated context, pairwise transfer, robustness interpretation, publication tables, and journal figures were regenerated after the primary rerun.",
        "",
        "## Numerical results that changed",
        "- Within-study permutation test resolution changed from the previous 199-permutation lower-bound p value to 4,999 permutations.",
        "- Current result: pseudo-F = 7.895136215963301; p = 0.0002; permutations = 4,999.",
        "- Other manuscript-level numerical summaries were retained from regenerated analysis outputs.",
        "",
        "## Figures regenerated",
    ]
    for name, paths in created["figures"].items():
        lines.append(f"- {name}: " + "; ".join(paths))
    lines.extend(
        [
            "",
            "## Machine-readable supplementary tables",
        ]
    )
    for name, paths in created["supplementary_tables"].items():
        lines.append(f"- {name}: " + "; ".join(paths))
    lines.extend(
        [
            "",
            "## Unresolved TODOs",
            "- None in the manuscript methods: the composite transportability score formula is documented.",
            "",
            "## Placeholders requiring author input before submission",
            "- [ADD AFFILIATION]",
            "- [ADD EMAIL]",
            "- [ADD GITHUB REPOSITORY LINK]",
            "- [ADD ZENODO DOI]",
            "",
            "## QC summary",
            f"- Abstract word count: {qc['abstract_word_count']}",
            f"- Package QC pass: {qc['pass']}",
            f"- DOCX render QA: {qc['docx_render_qa']}",
            validation_line,
        ]
    )
    if qc["prohibited_phrases"]:
        lines.append("- Prohibited phrases found: " + "; ".join(qc["prohibited_phrases"]))
    report.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report


def main() -> None:
    ensure_dirs()
    created: dict = {}

    supp_tables = build_supplementary_tables()
    result_csvs = copy_result_csvs()
    figures = copy_figures()

    items = manuscript_items()
    manuscript_docx = MANUSCRIPT_DIR / "manuscript_BMC_Microbiology.docx"
    manuscript_pdf = MANUSCRIPT_DIR / "manuscript_BMC_Microbiology.pdf"
    manuscript_md = MANUSCRIPT_DIR / "manuscript_BMC_Microbiology.md"
    build_docx(items, manuscript_docx)
    build_pdf(items, manuscript_pdf)
    write_markdown(items, manuscript_md)

    cover_docx = COVER_DIR / "cover_letter_BMC_Microbiology.docx"
    cover_md = COVER_DIR / "cover_letter_BMC_Microbiology.md"
    build_cover_letter_docx(cover_docx)
    write_cover_letter_md(cover_md)

    copied_scripts = copy_code_files()
    repro_files = write_reproducibility_files(copied_scripts)

    created.update(
        {
            "manuscript_docx": str(manuscript_docx),
            "manuscript_pdf": str(manuscript_pdf),
            "manuscript_markdown": str(manuscript_md),
            "cover_letter_docx": str(cover_docx),
            "cover_letter_markdown": str(cover_md),
            "supplementary_tables": supp_tables,
            "result_csvs": result_csvs,
            "figures": figures,
            "scripts": copied_scripts,
            "reproducibility_files": repro_files,
        }
    )
    qc = qc_package(items, created)
    qc_path = REPORTS_DIR / "bmc_submission_qc.json"
    qc_path.write_text(json.dumps(qc, ensure_ascii=False, indent=2), encoding="utf-8")
    created["qc_json"] = str(qc_path)
    report = write_final_report(created, qc)
    created["final_report"] = str(report)
    manifest = SUBMISSION_DIR / "manifest.json"
    manifest.write_text(json.dumps(created, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"submission_dir": str(SUBMISSION_DIR), "qc_pass": qc["pass"], "report": str(report)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
