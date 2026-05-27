from __future__ import annotations

import re
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.oxml.ns import qn
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


RUN_DIR = Path(r"D:\1\projects\gut_microbiome_colorectal_cancer\runs\run_20260525_182306")
BMC = RUN_DIR / "BMC_submission"
MANUSCRIPT_DOCX = BMC / "manuscript" / "manuscript_BMC_Microbiology.docx"
MANUSCRIPT_PDF = BMC / "manuscript" / "manuscript_BMC_Microbiology.pdf"
MANUSCRIPT_MD = BMC / "manuscript" / "manuscript_BMC_Microbiology.md"
COVER_DOCX = BMC / "cover_letter" / "cover_letter_BMC_Microbiology.docx"
COVER_MD = BMC / "cover_letter" / "cover_letter_BMC_Microbiology.md"

TITLE = "Transferability-aware prioritization of colorectal cancer gut metagenome signals across public cohorts"

SPECIES_NAMES = [
    "Parvimonas micra",
    "Gemella morbillorum",
    "Peptostreptococcus stomatis",
    "Streptococcus salivarius",
    "Roseburia intestinalis",
    "Faecalibacterium prausnitzii",
    "Roseburia faecis",
    "Eubacterium eligens",
    "Ruthenibacterium lactatiformans",
    "Bacteroides fragilis",
    "Butyricimonas virosa",
    "Flavonifractor plautii",
    "Anaerostipes hadrus",
    "Intestinimonas butyriciproducens",
    "Fusicatenibacter saccharivorans",
    "Parabacteroides distasonis",
    "Escherichia coli",
    "Blautia wexlerae",
    "Gemmiger formicilis",
    "Bilophila wadsworthia",
    "Bacteroides caccae",
    "Bacteroides thetaiotaomicron",
    "Clostridium bolteae",
    "Alistipes shahii",
    "Bifidobacterium longum",
    "Bifidobacterium adolescentis",
    "Roseburia inulinivorans",
    "Akkermansia muciniphila",
    "Agathobaculum butyriciproducens",
    "Collinsella stercoris",
    "Alistipes finegoldii",
    "Parabacteroides merdae",
]
SPECIES_RE = re.compile("(" + "|".join(re.escape(s) for s in sorted(SPECIES_NAMES, key=len, reverse=True)) + ")")


def split_species_runs(text: str):
    last = 0
    for match in SPECIES_RE.finditer(text):
        if match.start() > last:
            yield text[last : match.start()], False
        yield match.group(0), True
        last = match.end()
    if last < len(text):
        yield text[last:], False


def set_paragraph_text(paragraph, text: str, label: str | None = None) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    if label and text.startswith(f"{label}:"):
        run = paragraph.add_run(f"{label}: ")
        run.bold = True
        text = text[len(label) + 2 :].lstrip()
    for part, is_species in split_species_runs(text):
        run = paragraph.add_run(part)
        run.italic = is_species


def find_para(doc: Document, predicate) -> int:
    for i, p in enumerate(doc.paragraphs):
        if predicate(p):
            return i
    raise ValueError("paragraph not found")


def insert_before(paragraph, items: list[tuple[str, str]]) -> None:
    for style, text in reversed(items):
        new_p = paragraph.insert_paragraph_before()
        if style:
            new_p.style = style
        set_paragraph_text(new_p, text)


def revise_manuscript() -> None:
    doc = Document(MANUSCRIPT_DOCX)
    set_paragraph_text(doc.paragraphs[0], TITLE)

    replacements = {
        "Background:": "Background: Public processed metagenomic resources enable sample-level reanalysis of colorectal cancer (CRC)-associated gut microbiome signals across cohorts without starting from raw sequencing files. We evaluated whether transferability-aware candidate prioritization could separate association strength, ecological coherence, and cross-cohort portability in a public processed species matrix.",
        "Results:": "Results: We analyzed 1,395 CRC/control samples from 11 study labels and retained 301 species-level features. Variation partitioning attributed 8.47% of retained CLR variation to study label and 0.52% to CRC status after study adjustment. A study-adjusted association screen identified 124 species at FDR < 0.10, whereas 29 candidates met the stricter ≥60% same-direction cohort-support rule and 18 met high-consistency criteria. Candidate panel benchmarking showed that the top 29 species ranked by study-adjusted q value had the highest separability (LOSO median AUROC 0.834; pairwise off-diagonal median AUROC 0.733), while the high-consistency 18-species panel had lower transferability loss (0.186) than random 29-species panels (median 0.297). Ecological guild scores also diverged: butyrate/SCFA-associated commensal scores were control-higher (random-effects q = 1.0e-05), oral/pathobiont-associated scores were CRC-higher (q = 5.4e-05), and the Bacteroides/Enterobacteriaceae context score showed higher heterogeneity.",
        "Conclusions:": "Conclusions: Association strength, ecological coherence, and cohort transferability diverge in public CRC gut metagenomes. Transferability-aware prioritization provides a more conservative and interpretable candidate shortlist than pooled association ranking alone, but it does not establish causal, clinical diagnostic, stage-specific, mechanistic, functional, strain-level, oral-source, or raw-read reprocessing claims.",
    }
    for p in doc.paragraphs:
        for prefix, text in replacements.items():
            if p.text.startswith(prefix):
                set_paragraph_text(p, text, prefix[:-1])

    objective_i = find_para(doc, lambda p: p.text.startswith("The objective was to evaluate whether a sample-level public processed CRC/control metagenomic matrix"))
    set_paragraph_text(
        doc.paragraphs[objective_i],
        "The objective was to evaluate whether a sample-level public processed CRC/control metagenomic matrix can support transferability-aware candidate prioritization. The workflow was designed to separate study-adjusted association strength, ecological coherence, random-effects heterogeneity, cross-cohort direction stability, leave-one-study-out separability, pairwise study transfer, and sensitivity to analytic choices. This design differs from literature-level pooled analyses because it operates on harmonized sample-level processed abundance tables rather than published summary statistics. The analysis was therefore framed as a processed-data robustness and transferability audit, not as raw-read reprocessing, clinical diagnostic validation, or mechanistic inference.",
    )

    ai_i = find_para(doc, lambda p: p.style.name == "Heading 2" and p.text == "Use of AI-assisted tools")
    insert_before(
        doc.paragraphs[ai_i],
        [
            ("Heading 2", "Variance partitioning"),
            ("Normal", "To quantify the relative contribution of study label, CRC/control status, and available covariates, retained CLR profiles were analyzed with distance-equivalent multivariate linear models. Total Aitchison-space variation was calculated from centered retained CLR values. Partial variation explained was estimated from the reduction in residual sums of squares between reduced and full multivariate models. Study label was evaluated marginally; CRC status was evaluated after study label; age, BMI, and gender were evaluated in complete-case samples after study label, CRC status, and the other available covariates. These estimates were used as descriptive variance-partitioning metrics, not as causal attribution."),
            ("Heading 2", "Candidate panel benchmarking"),
            ("Normal", "Candidate panels were benchmarked under the same cross-cohort separability framework: all 301 retained species, the top 29 species ranked by study-adjusted BH q value, the 29 stable candidates, the 18 high-consistency candidates, and 500 random 29-species panels sampled without replacement from the retained feature set. Each panel was evaluated with leave-one-study-out and pairwise train-study × test-study elastic-net stress tests using training-only standardization and the same fixed model settings described above. Reported metrics included median AUROC, minimum AUROC, interquartile range, and transferability loss, defined as median within-study pairwise AUROC minus median off-diagonal pairwise AUROC. Candidate panels were selected from the processed matrix for prioritization stress testing and were not treated as nested clinical model-development pipelines."),
            ("Heading 2", "Ecological guild score analyses"),
            ("Normal", "Three transparent taxonomy-name guild panels were constructed before interpretation: a butyrate/SCFA-associated commensal panel, an oral/pathobiont-associated panel, and a Bacteroides/Enterobacteriaceae context panel. Guild scores were defined as the sample-wise mean CLR value across retained taxa in each panel. For each guild score, study-adjusted association, per-cohort effects, DerSimonian-Laird random-effects summaries, leave-one-cohort-out robustness, and LOSO/pairwise transferability stress tests were computed. Guild membership is reported exactly in Supplementary Table S10. These guild scores are ecological context summaries; they are not functional profiles, strain-level analyses, inflammation assays, or oral-source attribution models."),
        ],
    )

    h = find_para(doc, lambda p: p.style.name == "Heading 2" and p.text == "Processed matrix structure and compositional separation")
    set_paragraph_text(doc.paragraphs[h], "Processed matrix structure, variance partitioning, and compositional separation")
    set_paragraph_text(
        doc.paragraphs[h + 1],
        "The analysis covered 1,395 samples, including 701 CRC and 694 control records across 11 study labels. After prevalence and cohort-presence filtering, 301 of 944 species-level features were retained for CLR-based analysis. Variance partitioning showed that study label explained 8.47% of retained CLR variation, whereas CRC status explained 0.52% after study adjustment. In complete-case covariate models, age, gender, and BMI explained 0.37%, 0.21%, and 0.15%, respectively, after study label, CRC status, and the other available covariates. Aitchison PCA showed broad CRC/control overlap, but the within-study permutation test detected residualized CRC/control compositional separation (pseudo-F 7.90, p = 0.0002, 4,999 permutations; Figure 1; Figure 7A).",
    )

    ecology_h = find_para(doc, lambda p: p.style.name == "Heading 2" and p.text == "Ecological context, oral-associated signal, and benchmark context")
    insert_before(
        doc.paragraphs[ecology_h],
        [
            ("Heading 2", "Transferability-aware candidate panel benchmarking"),
            ("Normal", "Candidate panel benchmarking showed that association-ranked and transferability-aware panels behaved differently under cohort shift. The top 29 species ranked by study-adjusted BH q value had the highest separability, with LOSO median AUROC 0.834 and pairwise off-diagonal median AUROC 0.733. The 29 stable candidates had lower separability (LOSO median AUROC 0.673; pairwise off-diagonal median AUROC 0.595), indicating that same-direction cohort support does not automatically maximize separability. The 18 high-consistency candidates had LOSO median AUROC 0.693 and pairwise off-diagonal median AUROC 0.634, with transferability loss 0.186, lower than the random 29-species panel median loss of 0.297. These results support transferability-aware prioritization as a conservative interpretation layer rather than a route to maximizing apparent AUROC (Figure 7B-C; Supplementary Table S7)."),
        ],
    )
    set_paragraph_text(doc.paragraphs[ecology_h + 2], "Ecological guild scores and benchmark context")
    set_paragraph_text(
        doc.paragraphs[ecology_h + 3],
        "Ecological guild score analyses separated coherent ecological summaries from individual-species ranking. The butyrate/SCFA-associated commensal score was control-higher in the study-adjusted model (coefficient -0.468, q = 8.59e-10) and remained supported in the random-effects audit (random-effects coefficient -0.503, q = 1.0e-05, I² = 51.5%). The oral/pathobiont-associated score was CRC-higher (study-adjusted coefficient 0.206, q = 6.75e-05; random-effects coefficient 0.193, q = 5.4e-05, I² = 5.1%). The Bacteroides/Enterobacteriaceae context score was CRC-higher in the study-adjusted model (coefficient 0.109, q = 0.0196), but the random-effects summary was not significant at FDR < 0.10 and showed higher heterogeneity (q = 0.231, I² = 70.3%). All three guild scores met the predefined leave-one-cohort-out direction-preservation criterion (Figure 7D; Supplementary Tables S9-S10).",
    )
    set_paragraph_text(
        doc.paragraphs[ecology_h + 4],
        "The stable-candidate list and MicrobiomeHD benchmark context were related but not identical. Among stable candidates, 17 had genus-level MicrobiomeHD context, and one belonged to the literature-prior oral-associated panel. This comparison provides external processed-table context but does not merge individual samples, harmonize schemas, or validate a classifier. It was therefore interpreted as benchmark context rather than external validation.",
    )

    discussion_first = find_para(doc, lambda p: p.text.startswith("This sample-level public processed-matrix reanalysis shows"))
    set_paragraph_text(
        doc.paragraphs[discussion_first],
        "This sample-level public processed-matrix reanalysis shows that association strength, ecological coherence, and cohort transferability diverge in public CRC gut metagenomes. Study label explained substantially more retained CLR variation than CRC status, and candidate panel benchmarking showed that the q-ranked top 29 panel maximized separability while the high-consistency panel provided a more conservative shortlist with lower transferability loss than random panels. The most important result is therefore not a single taxon, guild score, or AUROC value; it is the observation that prioritization changes when association ranking, ecological context, heterogeneity, and transferability are evaluated as separate evidence layers.",
    )
    discussion_final = find_para(doc, lambda p: p.text.startswith("Strengths of this study include explicit study adjustment"))
    set_paragraph_text(
        doc.paragraphs[discussion_final],
        "A strength of the revised workflow is that it tests candidate lists against cohort shift rather than relying only on pooled association ranking. Variance partitioning, panel benchmarking, random-panel comparisons, guild-score summaries, heterogeneity auditing, LOCO sensitivity, LOSO stress testing, and pairwise study transfer each answer a different question. The limitations remain important: the analysis used public processed tables only; candidate panels were selected from the available processed matrix; no raw reads, functional profiles, strains, disease-stage models, independent patient-level validation matrices, or clinical calibration were analyzed. The resulting shortlist is therefore conservative and interpretable for processed-data candidate prioritization, not causal, clinical diagnostic, mechanistic, oral-source, or raw-read evidence.",
    )

    conc1 = find_para(doc, lambda p: p.text.startswith("Public processed CRC/control gut metagenome matrices can support more than a simple association screen"))
    set_paragraph_text(
        doc.paragraphs[conc1],
        "Public processed CRC/control gut metagenome matrices can support transferability-aware candidate prioritization when association strength, ecological coherence, heterogeneity, and cohort portability are evaluated separately. In 1,395 samples across 11 study labels, study label explained 8.47% of retained CLR variation and CRC status explained 0.52% after study adjustment. The retained 301-species matrix produced 124 study-adjusted associations at FDR < 0.10, 29 same-direction stable candidates under a ≥60% cohort-support rule, and 18 high-consistency candidates.",
    )
    set_paragraph_text(
        doc.paragraphs[conc1 + 1],
        "The revised analysis shows that pooled association ranking, ecological guild behavior, and transferability are not interchangeable. The q-ranked top 29 panel had the highest cross-cohort separability, whereas the high-consistency panel was more conservative and had lower transferability loss than random 29-species panels. Butyrate/SCFA-associated commensal depletion and oral/pathobiont-associated enrichment were coherent at the guild-score level, while the Bacteroides/Enterobacteriaceae context score was more heterogeneous. These findings support a reproducible processed-data candidate-prioritization workflow, but they do not establish causal, clinical diagnostic, stage-specific, mechanistic, functional, strain-level, oral-source, or raw-read reprocessing claims.",
    )

    refs_i = find_para(doc, lambda p: p.style.name == "Heading 1" and p.text == "References")
    additional = [
        ("Normal", "Additional file 8: Supplementary Table S6. Variance partitioning of retained CLR profiles."),
        ("Normal", "Additional file 9: Supplementary Table S7. Candidate panel benchmarking results."),
        ("Normal", "Additional file 10: Supplementary Table S8. Candidate panel membership."),
        ("Normal", "Additional file 11: Supplementary Table S9. Ecological guild score results."),
        ("Normal", "Additional file 12: Supplementary Table S10. Ecological guild panel membership."),
    ]
    insert_before(doc.paragraphs[refs_i], additional)

    fig_legend_end = find_para(doc, lambda p: p.text.startswith("Supplementary Figure S1."))
    insert_before(
        doc.paragraphs[fig_legend_end],
        [
            ("Normal", "Figure 7. Transferability-aware candidate prioritization. (A) Distance-equivalent variance partitioning of retained CLR profiles by study label, CRC status, and available covariates. (B) LOSO and pairwise off-diagonal AUROC summaries across fixed candidate panels and random 29-species panels. (C) Fixed panels overlaid on the distribution from 500 random 29-species panels for LOSO median AUROC, pairwise off-diagonal median AUROC, and transferability loss. (D) Random-effects summaries for ecological guild scores. These analyses support processed-data prioritization and do not represent clinical diagnostic validation."),
        ],
    )

    doc.save(MANUSCRIPT_DOCX)


def revise_cover_letter() -> None:
    doc = Document(COVER_DOCX)
    for p in doc.paragraphs:
        txt = p.text
        if "Robustness and transferability of colorectal cancer gut metagenome signals across public cohorts" in txt:
            set_paragraph_text(p, txt.replace("Robustness and transferability of colorectal cancer gut metagenome signals across public cohorts", TITLE))
        elif txt.startswith("This study presents a reproducible cohort-aware computational reanalysis"):
            set_paragraph_text(
                p,
                "This study presents a reproducible cohort-aware computational reanalysis of sample-level public processed colorectal cancer gut metagenome matrices. The revised analysis adds a transferability-aware candidate prioritization layer that separates association strength, ecological guild coherence, random-effects heterogeneity, cross-cohort direction stability, leave-one-study-out separability, pairwise cross-study transfer, random-panel benchmarking, and sensitivity to feature filtering and pseudocount choice.",
            )
        elif txt.startswith("The manuscript is aligned with BMC Microbiology"):
            set_paragraph_text(
                p,
                "The manuscript is aligned with BMC Microbiology because it addresses human gut microbiome signals using a transparent and reproducible computational workflow. The study is not a literature-level pooled analysis, bibliometric analysis, raw-read reprocessing study, or clinical diagnostic validation study. Instead, it provides a bounded processed-data candidate-prioritization workflow and shows that association strength, ecological coherence, and cohort transferability can select partially different microbial signals.",
            )
    doc.save(COVER_DOCX)
    COVER_MD.write_text("\n\n".join(p.text for p in doc.paragraphs if p.text.strip()) + "\n", encoding="utf-8")


def register_pdf_fonts() -> tuple[str, str]:
    regular = Path(r"C:\Windows\Fonts\arial.ttf")
    bold = Path(r"C:\Windows\Fonts\arialbd.ttf")
    if regular.exists() and bold.exists():
        pdfmetrics.registerFont(TTFont("Arial", str(regular)))
        pdfmetrics.registerFont(TTFont("Arial-Bold", str(bold)))
        return "Arial", "Arial-Bold"
    return "Times-Roman", "Times-Bold"


def html_species(text: str) -> str:
    escaped = escape(text)
    for species in sorted(SPECIES_NAMES, key=len, reverse=True):
        escaped = escaped.replace(escape(species), f"<i>{escape(species)}</i>")
    return escaped


def page_number(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Times-Roman", 9)
    canvas.drawCentredString(4.25 * inch, 0.45 * inch, f"Page {doc.page}")
    canvas.restoreState()


def regenerate_md_and_pdf() -> None:
    doc = Document(MANUSCRIPT_DOCX)
    lines = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        if p.style.name == "Title":
            lines.append(f"# {p.text}")
        elif p.style.name == "Heading 1":
            lines.append(f"\n## {p.text}")
        elif p.style.name == "Heading 2":
            lines.append(f"\n### {p.text}")
        else:
            lines.append(f"\n{p.text}")
    MANUSCRIPT_MD.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")

    regular, bold = register_pdf_fonts()
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName=regular, fontSize=10.5, leading=18, spaceAfter=6)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName=bold, fontSize=14, leading=20, spaceBefore=10, spaceAfter=6)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName=bold, fontSize=12, leading=18, spaceBefore=8, spaceAfter=4)
    title = ParagraphStyle("Title", parent=styles["Title"], fontName=bold, fontSize=16, leading=22, alignment=TA_CENTER, spaceAfter=12)
    story = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name == "Title":
            story.append(Paragraph(html_species(text), title))
        elif p.style.name == "Heading 1":
            story.append(Paragraph(html_species(text), h1))
        elif p.style.name == "Heading 2":
            story.append(Paragraph(html_species(text), h2))
        else:
            story.append(Paragraph(html_species(text), body))
        story.append(Spacer(1, 0.03 * inch))
    pdf = SimpleDocTemplate(str(MANUSCRIPT_PDF), pagesize=LETTER, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    pdf.build(story, onFirstPage=page_number, onLaterPages=page_number)


def main() -> None:
    revise_manuscript()
    revise_cover_letter()
    regenerate_md_and_pdf()
    print(
        {
            "manuscript_docx": str(MANUSCRIPT_DOCX),
            "manuscript_pdf": str(MANUSCRIPT_PDF),
            "cover_letter_docx": str(COVER_DOCX),
            "title": TITLE,
        }
    )


if __name__ == "__main__":
    main()
